"""生成 P0-3 至 P0-9 的隔离 Word 集成测试文稿和机器可读结果。"""
from __future__ import annotations

import argparse
import json
from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree

from docmask.core.codebook import Codebook
from docmask.core.masker import Masker
from docmask.core.restorer import Restorer
from docmask.handlers.docx_handler import DocxHandler, NSMAP
from tests.word_fixture_factory import create_complex_word_fixture


def _root_word_snapshot(project_root: Path) -> list[dict]:
    snapshot = []
    for path in sorted(project_root.glob("*.doc*"), key=lambda item: item.name.lower()):
        stat = path.stat()
        snapshot.append({
            "name": path.name,
            "size": stat.st_size,
            "mtime_ns": stat.st_mtime_ns,
        })
    return snapshot


def _scan_package(path: Path, original: str, replacement: str) -> dict:
    original_count = 0
    replacement_count = 0
    members = []
    with ZipFile(path) as package:
        corrupt_member = package.testzip()
        for member in package.namelist():
            if member.startswith("word/") and member.endswith(".xml"):
                content = package.read(member)
                try:
                    member_root = etree.fromstring(content)
                except etree.XMLSyntaxError:
                    continue
                visible_text = "".join(
                    node.text or ""
                    for node in member_root.findall(".//w:t", NSMAP)
                )
                original_count += visible_text.count(original)
                replacement_count += visible_text.count(replacement)
                members.append(member)
        document_root = etree.fromstring(package.read("word/document.xml"))

    return {
        "zip_valid": corrupt_member is None,
        "word_xml_members": len(members),
        "original_count": original_count,
        "replacement_count": replacement_count,
        "structure": {
            "tabs": len(document_root.findall(".//w:tab", NSMAP)),
            "breaks": len(document_root.findall(".//w:br", NSMAP)),
            "drawings": len(document_root.findall(".//w:drawing", NSMAP)),
            "field_chars": len(document_root.findall(".//w:fldChar", NSMAP)),
            "field_instructions": len(document_root.findall(".//w:instrText", NSMAP)),
            "hyperlinks": len(document_root.findall(".//w:hyperlink", NSMAP)),
        },
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--project-root", required=True)
    parser.add_argument("--workspace", required=True)
    args = parser.parse_args()

    project_root = Path(args.project_root).resolve()
    workspace = Path(args.workspace).resolve()
    input_dir = workspace / "input"
    output_dir = workspace / "output"
    config_dir = workspace / "config"
    for directory in (input_dir, output_dir, config_dir):
        directory.mkdir(parents=True, exist_ok=False)

    before_snapshot = _root_word_snapshot(project_root)

    source_path = create_complex_word_fixture(input_dir / "complex_source.docx")
    codebook_path = config_dir / "codebook.txt"
    codebook_path.write_text("张三==>匿名客户001\n", encoding="utf-8")
    codebook = Codebook(str(codebook_path))
    codebook.load()

    handler = DocxHandler()
    masked_path, masked_count, coverage = handler.mask(
        str(source_path),
        Masker(codebook),
        output_path=str(output_dir / "complex_masked.docx"),
    )
    restored_path, restored_count = handler.restore(
        masked_path,
        Restorer(codebook),
        output_path=str(output_dir / "complex_restored.docx"),
    )

    source_scan = _scan_package(source_path, "张三", "匿名客户001")
    masked_scan = _scan_package(Path(masked_path), "张三", "匿名客户001")
    restored_scan = _scan_package(Path(restored_path), "张三", "匿名客户001")

    masked_document = Document(masked_path)
    masked_hyperlinks = [
        hyperlink.text
        for paragraph in masked_document.paragraphs
        for hyperlink in paragraph.hyperlinks
    ]
    restored_document = Document(restored_path)
    restored_hyperlinks = [
        hyperlink.text
        for paragraph in restored_document.paragraphs
        for hyperlink in paragraph.hyperlinks
    ]

    after_snapshot = _root_word_snapshot(project_root)
    checks = {
        "existing_root_word_documents_unchanged": before_snapshot == after_snapshot,
        "source_contains_original": source_scan["original_count"] > 0,
        "masked_has_no_original": masked_scan["original_count"] == 0,
        "masked_contains_replacement": masked_scan["replacement_count"] > 0,
        "restored_has_no_replacement": restored_scan["replacement_count"] == 0,
        "restored_original_count_matches_source": (
            restored_scan["original_count"] == source_scan["original_count"]
        ),
        "structure_preserved_after_mask": (
            masked_scan["structure"] == source_scan["structure"]
        ),
        "structure_preserved_after_restore": (
            restored_scan["structure"] == source_scan["structure"]
        ),
        "word_xml_part_count_preserved_after_mask": (
            masked_scan["word_xml_members"] == source_scan["word_xml_members"]
        ),
        "word_xml_part_count_preserved_after_restore": (
            restored_scan["word_xml_members"] == source_scan["word_xml_members"]
        ),
        "all_packages_are_valid_zip": all(
            scan["zip_valid"] for scan in (source_scan, masked_scan, restored_scan)
        ),
        "hyperlink_masked": "匿名客户001" in masked_hyperlinks,
        "hyperlink_restored": "张三" in restored_hyperlinks,
        "replacement_counts_round_trip": restored_count == masked_count,
    }

    result = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "project_root": str(project_root),
        "workspace": str(workspace),
        "isolation": {
            "existing_word_documents_before": before_snapshot,
            "existing_word_documents_after": after_snapshot,
        },
        "artifacts": {
            "source": str(source_path),
            "masked": str(masked_path),
            "restored": str(restored_path),
            "codebook": str(codebook_path),
        },
        "mask_replacements": masked_count,
        "restore_replacements": restored_count,
        "coverage": coverage,
        "source_scan": source_scan,
        "masked_scan": masked_scan,
        "restored_scan": restored_scan,
        "checks": checks,
        "passed": all(checks.values()),
    }
    (workspace / "integration-result.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if result["passed"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
