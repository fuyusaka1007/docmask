"""P0-10～P0-12、P1-1～P1-6 审计项回归测试。"""
import base64
import json
import zipfile
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches
from lxml import etree

from docmask.core.codebook import Codebook, CodebookError
from docmask.core.masker import Masker
from docmask.core.restorer import Restorer
from docmask.handlers.docx_handler import DocxHandler
from docmask.handlers.txt_handler import TxtHandler
from docmask.services.file_service import scan_files
from docmask.ui.state import AppState, create_file_item
from docmask.ui.pages.results_page import ResultsPage
from docmask.utils.file_utils import resolve_output_path


def make_codebook(tmp_path: Path, rules: str) -> Codebook:
    path = tmp_path / "codebook.txt"
    path.write_text(rules, encoding="utf-8")
    codebook = Codebook(str(path))
    codebook.load()
    return codebook


def test_empty_matching_regex_is_rejected(tmp_path):
    path = tmp_path / "codebook.txt"
    path.write_text("regex:.*?==>X\n", encoding="utf-8")
    with pytest.raises(CodebookError, match="可匹配空字符串"):
        Codebook(str(path)).load()


def test_replacements_are_globally_unique_across_rule_types(tmp_path):
    codebook = make_codebook(tmp_path, "A==>MASK\nregex:\\d+==>MASK\n")
    assert any(message.startswith("ERROR") for message in codebook.validate())


def test_regex_replacement_is_literal_and_generated_text_is_not_reprocessed(tmp_path):
    codebook = make_codebook(
        tmp_path, "A==>123\nregex:\\d+==>\\1-LITERAL\n"
    )
    result, count, _hits = Masker(codebook).mask_text("A 456")
    assert result == "123 \\1-LITERAL"
    assert count == 2


def test_document_reversibility_is_separate_from_structure_validation(tmp_path):
    codebook = make_codebook(tmp_path, "A==>MASK\nregex:\\d+==>NUMBER\n")
    assert not [m for m in codebook.validate() if m.startswith("ERROR")]
    messages = codebook.validate_reversibility("A 123")
    assert any("正则规则" in message and "不能保证" in message for message in messages)


def test_cli_output_resolution_distinguishes_single_and_batch(tmp_path):
    source = tmp_path / "source.txt"
    source.write_text("x", encoding="utf-8")
    explicit = tmp_path / "named.txt"
    explicit.write_text("keep", encoding="utf-8")

    single = resolve_output_path(
        str(source), str(explicit), suffix="_masked", batch_mode=False
    )
    assert Path(single).name == "named_1.txt"
    assert explicit.read_text(encoding="utf-8") == "keep"

    output_dir = tmp_path / "batch-output"
    batch = resolve_output_path(
        str(source), str(output_dir), suffix="_masked", batch_mode=True
    )
    assert Path(batch) == output_dir / "source_masked.txt"


def test_handler_refuses_to_overwrite_explicit_output(tmp_path):
    source = tmp_path / "source.txt"
    source.write_text("A", encoding="utf-8")
    output = tmp_path / "output.txt"
    output.write_text("keep", encoding="utf-8")
    codebook = make_codebook(tmp_path, "A==>B\n")

    with pytest.raises(FileExistsError):
        TxtHandler().mask(str(source), Masker(codebook), output_path=str(output))
    assert output.read_text(encoding="utf-8") == "keep"


def test_docx_nested_table_and_tracked_insertion_round_trip(tmp_path):
    source = tmp_path / "complex.docx"
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).paragraphs[0].add_run("SECRET")

    paragraph = doc.add_paragraph()
    insertion = OxmlElement("w:ins")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "SECRET"
    run.append(text)
    insertion.append(run)
    paragraph._p.append(insertion)
    doc.save(source)

    codebook = make_codebook(tmp_path, "SECRET==>MASKED\n")
    handler = DocxHandler()
    masked, masked_count, _ = handler.mask(str(source), Masker(codebook))
    restored, restored_count = handler.restore(masked, Restorer(codebook))

    masked_xml = Document(masked).element.xml
    restored_xml = Document(restored).element.xml
    assert masked_count == 2
    assert restored_count == 2
    assert "SECRET" not in masked_xml
    assert masked_xml.count("MASKED") == 2
    assert restored_xml.count("SECRET") == 2


def test_docx_declares_capability_boundaries():
    matrix = DocxHandler.OPC_CAPABILITY_MATRIX
    assert matrix["word/comments*.xml"] == "supported"
    assert matrix["word/charts/*"] == "warn-only"
    assert matrix["external hyperlink targets"] == "blocked-by-default"
    # A-01: 扩展隐私表面声明
    surface = DocxHandler.EXTENDED_PRIVACY_SURFACE
    assert "instrText" in surface["text_nodes"]
    assert "tag" in surface["attributes"]
    assert "bookmarkStart" in surface["attributes"]
    assert "docPr" in surface["attributes"]


def test_docx_comment_text_and_author_round_trip(tmp_path):
    """批注文本和作者脱敏/恢复往返测试。

    注意：python-docx 1.1.0 无 Document.comments/add_comment API，
    但 DocxHandler 在 XML 层级处理 word/comments.xml，因此测试直接
    操作 OPC 包来创建和验证批注。
    """
    from lxml import etree
    from io import BytesIO
    import zipfile
    import re

    source = tmp_path / "comments.docx"
    doc = Document()
    doc.add_paragraph("anchor text")
    # 手动注入 word/comments.xml 到 OPC 包
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    comments_xml = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        b'<w:comment w:id="0" w:author="SECRET" w:initials="S">'
        b'<w:p><w:r><w:t>SECRET comment</w:t></w:r></w:p>'
        b'</w:comment></w:comments>'
    )

    with (
        zipfile.ZipFile(buf, "r") as zin,
        zipfile.ZipFile(source, "w", zipfile.ZIP_DEFLATED) as zout,
    ):
        # 复制所有现有条目
        for item in zin.infolist():
            if item.filename == "[Content_Types].xml":
                data = zin.read(item)
                # 注入 comments 的内容类型声明
                before = b'</Types>'
                override = b'<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
                data = data.replace(before, override + before)
                zout.writestr(item, data)
            elif item.filename == "word/_rels/document.xml.rels":
                data = zin.read(item)
                before = b'</Relationships>'
                rel = (
                    b'<Relationship Id="rIdComments" '
                    b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" '
                    b'Target="comments.xml"/>'
                )
                data = data.replace(before, rel + before)
                zout.writestr(item, data)
            else:
                zout.writestr(item, zin.read(item))
        # 写入 comments.xml
        zout.writestr("word/comments.xml", comments_xml)

    codebook = make_codebook(tmp_path, "SECRET==>MASKED\n")
    handler = DocxHandler()

    masked, count, _ = handler.mask(str(source), Masker(codebook))
    restored, restored_count = handler.restore(masked, Restorer(codebook))

    # 通过 XML 直接验证（不走 Document.comments）
    def get_comment_info(docx_path):
        with zipfile.ZipFile(docx_path, "r") as z:
            xml_bytes = z.read("word/comments.xml")
        root = etree.fromstring(xml_bytes)
        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        comment = root.find(".//w:comment", nsmap)
        text_nodes = [t.text for t in comment.findall(".//w:t", nsmap) if t.text]
        return {
            "author": comment.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author"
            ),
            "initials": comment.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}initials"
            ),
            "text": "".join(text_nodes),
        }

    masked_info = get_comment_info(masked)
    restored_info = get_comment_info(restored)

    assert count == 2  # "SECRET" in text + author
    assert masked_info["text"] == "MASKED comment"
    assert masked_info["author"] == "MASKED"
    assert restored_count == 2
    assert restored_info["text"] == "SECRET comment"
    assert restored_info["author"] == "SECRET"


def test_docx_unsupported_text_part_emits_warning_without_leaking_value(tmp_path):
    codebook = make_codebook(tmp_path, "SECRET==>MASKED\n")
    chart_xml = (
        b'<c:chart xmlns:c="urn:chart" xmlns:a="urn:drawing">'
        b'<a:t>SECRET</a:t></c:chart>'
    )
    part = SimpleNamespace(
        partname="/word/charts/chart1.xml", blob=chart_xml, rels={}
    )
    doc = SimpleNamespace(
        part=SimpleNamespace(package=SimpleNamespace(parts=[part]))
    )
    warnings = DocxHandler()._scan_unsupported_parts(doc, Masker(codebook), original_texts={"SECRET"})
    assert warnings == ["不支持的 OPC 文本部件可能含敏感内容：word/charts/chart1.xml"]
    assert "SECRET" not in warnings[0]


def test_single_pass_folder_scan_returns_skipped_and_progress(tmp_path):
    (tmp_path / "a.txt").write_text("a", encoding="utf-8")
    (tmp_path / "b.pdf").write_text("b", encoding="utf-8")
    progress = []
    files, skipped, errors = scan_files(
        str(tmp_path), ["txt"], progress_callback=lambda count, path: progress.append((count, path))
    )
    assert [Path(path).name for path in files] == ["a.txt"]
    assert skipped == 1
    assert errors == []
    assert progress


def test_custom_output_directory_is_execution_precondition(tmp_path):
    state = AppState()
    state.codebook.valid = True
    source = tmp_path / "a.txt"
    source.write_text("a", encoding="utf-8")
    state.files = [create_file_item(str(source))]
    state.output_same_dir = False
    state.output_dir = None
    assert state.can_execute is False
    state.output_dir = str(tmp_path)
    assert state.can_execute is True


def test_results_empty_state_resets_all_cards(monkeypatch):
    class Card:
        def __init__(self):
            self.value = "stale"

        def set_value(self, value):
            self.value = value

    class Scroll:
        def winfo_children(self):
            return []

        content = SimpleNamespace(winfo_children=lambda: [])

    class Label:
        def __init__(self, *_args, **_kwargs):
            pass

        def pack(self, **_kwargs):
            pass

    monkeypatch.setattr("docmask.ui.pages.results_page.ctk.CTkLabel", Label)
    page = SimpleNamespace(
        state=AppState(),
        _list_scroll=Scroll(),
        _card_total=Card(),
        _card_success=Card(),
        _card_fail=Card(),
        _card_stopped=Card(),
        _card_replacements=Card(),
    )
    ResultsPage._refresh(page)
    assert {
        page._card_total.value,
        page._card_success.value,
        page._card_fail.value,
        page._card_stopped.value,
        page._card_replacements.value,
    } == {"0"}


# ===== A-01: 扩展隐私表面脱敏/恢复/校验 =====

_ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _build_extended_privacy_fixture(path: Path) -> Path:
    """创建包含域代码/SDT标签/书签名/图片替代文本的 DOCX。"""
    doc = Document()

    # 1. w:instrText 域代码指令
    p1 = doc.add_paragraph()
    r1 = p1.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = " MERGEFIELD SECRET_FIELD "
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t"); display.text = "placeholder"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, separate, display, end):
        r1._r.append(el)

    # 2. w:sdtPr/w:tag/@w:val 内容控件标签
    p2 = doc.add_paragraph()
    sdt = OxmlElement("w:sdt")
    sdtPr = OxmlElement("w:sdtPr")
    tag = OxmlElement("w:tag"); tag.set(qn("w:val"), "SECRET_TAG")
    sdtPr.append(tag)
    sdtContent = OxmlElement("w:sdtContent")
    sdt_p = OxmlElement("w:p")
    sdt_r = OxmlElement("w:r")
    sdt_t = OxmlElement("w:t"); sdt_t.text = "content"
    sdt_r.append(sdt_t); sdt_p.append(sdt_r)
    sdtContent.append(sdt_p)
    sdt.append(sdtPr); sdt.append(sdtContent)
    p2._p.append(sdt)

    # 3. w:bookmarkStart/@w:name 书签名称
    p3 = doc.add_paragraph("bookmark target")
    bookmark = OxmlElement("w:bookmarkStart")
    bookmark.set(qn("w:id"), "0")
    bookmark.set(qn("w:name"), "SECRET_BOOKMARK")
    p3._p.insert(0, bookmark)
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), "0")
    p3._p.append(bookmark_end)

    # 4. wp:docPr/@descr @title 图片替代文本
    p4 = doc.add_paragraph()
    run4 = p4.add_run()
    run4.add_picture(BytesIO(_ONE_PIXEL_PNG), width=Inches(0.2))
    for elem in doc.element.body.iter():
        if etree.QName(elem).localname == "docPr":
            elem.set("descr", "SECRET_ALT")
            elem.set("title", "SECRET_ALT_TITLE")
            break

    doc.save(path)
    return path


def test_docx_extended_privacy_surface_masked_and_restored(tmp_path):
    """A-01: 域代码/SDT标签/书签名/图片替代文本必须被脱敏和恢复。"""
    source = _build_extended_privacy_fixture(tmp_path / "extended.docx")

    codebook = make_codebook(tmp_path, (
        "SECRET_FIELD==>MASKED_FIELD\n"
        "SECRET_TAG==>MASKED_TAG\n"
        "SECRET_BOOKMARK==>MASKED_BOOKMARK\n"
        "SECRET_ALT==>MASKED_ALT\n"
        "SECRET_ALT_TITLE==>MASKED_ALT_TITLE\n"
    ))
    handler = DocxHandler()
    masked, count, _ = handler.mask(str(source), Masker(codebook))

    assert count >= 5  # 至少 5 处替换

    with zipfile.ZipFile(masked) as z:
        document_xml = z.read("word/document.xml").decode("utf-8")

    # 脱敏后不应残留原文
    for secret in (
        "SECRET_FIELD", "SECRET_TAG", "SECRET_BOOKMARK",
        "SECRET_ALT", "SECRET_ALT_TITLE",
    ):
        assert secret not in document_xml, f"扩展隐私表面残留：{secret}"
    # 脱敏词应存在
    for masked_val in ("MASKED_FIELD", "MASKED_TAG", "MASKED_BOOKMARK", "MASKED_ALT"):
        assert masked_val in document_xml

    # 恢复往返
    restored, restore_count = handler.restore(masked, Restorer(codebook))
    with zipfile.ZipFile(restored) as z:
        restored_xml = z.read("word/document.xml").decode("utf-8")
    for secret in ("SECRET_FIELD", "SECRET_TAG", "SECRET_BOOKMARK", "SECRET_ALT"):
        assert secret in restored_xml, f"恢复后缺失：{secret}"


def test_docx_bookmark_name_invalid_replacement_blocks_output(tmp_path):
    """A-01: 书签名脱敏词不符合 XML Name 规范时 fail closed。"""
    source = tmp_path / "bookmark.docx"
    doc = Document()
    p = doc.add_paragraph("text")
    bookmark = OxmlElement("w:bookmarkStart")
    bookmark.set(qn("w:id"), "0")
    bookmark.set(qn("w:name"), "SECRET_BOOKMARK")
    p._p.insert(0, bookmark)
    doc.save(source)

    # {MASKED} 不是合法 XML Name（以 { 开头）
    codebook = make_codebook(tmp_path, "SECRET_BOOKMARK==>{MASKED}\n")
    handler = DocxHandler()
    with pytest.raises(RuntimeError, match="书签名称"):
        handler.mask(str(source), Masker(codebook))
    # 阻止输出：不应产生脱敏文件
    assert not list(tmp_path.glob("*_desensitized*"))


def test_docx_external_relationship_secret_blocked_by_default(tmp_path):
    """A-01: 外部关系目标含敏感内容时，安全模式默认阻止输出。"""
    source = tmp_path / "hyperlink.docx"
    doc = Document()
    p = doc.add_paragraph()
    rid = p.part.relate_to(
        "https://example.invalid/SECRET_URL", RT.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    t = OxmlElement("w:t"); t.text = "link text"
    run.append(t); hyperlink.append(run)
    p._p.append(hyperlink)
    doc.save(source)

    codebook = make_codebook(tmp_path, "SECRET_URL==>MASKED_URL\n")

    # 默认安全模式：阻止输出
    handler = DocxHandler()
    with pytest.raises(RuntimeError, match="外部关系目标含敏感内容"):
        handler.mask(str(source), Masker(codebook))
    assert not list(tmp_path.glob("*_desensitized*"))

    # 兼容模式：仅告警
    handler2 = DocxHandler()
    handler2.strict_external_targets = False
    masked, _, _ = handler2.mask(str(source), Masker(codebook))
    assert any("外部关系目标" in w for w in handler2.last_warnings)


# ===== A-03: 不支持部件跨节点告警 =====


def test_docx_unsupported_part_cross_node_split_emits_warning(tmp_path):
    """A-03: 敏感原文被拆分到相邻文本节点时，滑动窗口聚合扫描应告警。"""
    codebook = make_codebook(tmp_path, "SECRET==>MASKED\n")
    # SECRET 被拆成 SEC + RET 两个相邻 a:t 节点
    chart_xml = (
        b'<c:chart xmlns:c="urn:chart" xmlns:a="urn:drawing">'
        b'<a:t>SEC</a:t><a:t>RET</a:t></c:chart>'
    )
    part = SimpleNamespace(
        partname="/word/charts/chart1.xml", blob=chart_xml, rels={}
    )
    doc = SimpleNamespace(
        part=SimpleNamespace(package=SimpleNamespace(parts=[part]))
    )
    warnings = DocxHandler()._scan_unsupported_parts(
        doc, Masker(codebook), original_texts={"SECRET"}
    )
    assert len(warnings) == 1
    assert "跨节点" in warnings[0]
    assert "SECRET" not in warnings[0]


# ===== A-04: 无硬链接文件系统安全回退 =====


def test_staged_output_falls_back_when_hardlink_unavailable(tmp_path, monkeypatch):
    """A-04: os.link 失败时应回退到 O_CREAT|O_EXCL + 复制 + fsync。"""
    from docmask.utils.file_utils import staged_output_path

    # 模拟不支持硬链接的文件系统
    monkeypatch.setattr("os.link", lambda *a, **kw: (_ for _ in ()).throw(OSError("EPERM")))

    final = tmp_path / "output.txt"
    with staged_output_path(str(final)) as temp:
        Path(temp).write_text("测试内容", encoding="utf-8")

    assert final.exists()
    assert final.read_text(encoding="utf-8") == "测试内容"


def test_staged_output_refuses_overwrite_even_in_fallback(tmp_path, monkeypatch):
    """A-04: 回退路径也必须保证不覆盖已有文件。"""
    from docmask.utils.file_utils import staged_output_path

    monkeypatch.setattr("os.link", lambda *a, **kw: (_ for _ in ()).throw(OSError("EPERM")))

    final = tmp_path / "existing.txt"
    final.write_text("原有内容", encoding="utf-8")

    with pytest.raises(FileExistsError):
        with staged_output_path(str(final)) as temp:
            Path(temp).write_text("新内容", encoding="utf-8")

    # 原文件未被覆盖
    assert final.read_text(encoding="utf-8") == "原有内容"


# ===== A-05: 设置类型校验 =====


def test_settings_load_validates_bad_values_and_falls_back_per_field(tmp_path, monkeypatch):
    """A-05: 坏字段回退默认值，好字段保留。"""
    from docmask.ui.state import SettingsModel

    settings_file = tmp_path / "settings.json"
    settings_file.write_text(json.dumps({
        "theme": [],              # 无效 -> 回退默认
        "scale": "150%",          # 无效 -> 回退默认
        "log_level": "VERBOSE",   # 无效 -> 回退默认
        "format_filters": None,   # 无效 -> 回退默认
        "output_same_dir": "yes", # 无效 -> 回退默认
        "generate_report": False,  # 有效
    }), encoding="utf-8")

    monkeypatch.setattr(SettingsModel, "_settings_path", staticmethod(lambda: settings_file))
    model = SettingsModel.load()

    assert model.theme == "跟随系统"
    assert model.scale == "100%"
    assert model.log_level == "INFO"
    assert model.format_filters == ["docx", "doc", "txt"]
    assert model.output_same_dir is True
    assert model.generate_report is False  # 好字段保留


def test_settings_load_preserves_valid_values(tmp_path, monkeypatch):
    """A-05: 全部合法字段正常加载。"""
    from docmask.ui.state import SettingsModel

    settings_file = tmp_path / "settings.json"
    settings_file.write_text(json.dumps({
        "theme": "深色",
        "scale": "120%",
        "log_level": "DEBUG",
        "format_filters": ["docx", "txt"],
        "output_same_dir": False,
        "generate_report": False,
    }), encoding="utf-8")

    monkeypatch.setattr(SettingsModel, "_settings_path", staticmethod(lambda: settings_file))
    model = SettingsModel.load()

    assert model.theme == "深色"
    assert model.scale == "120%"
    assert model.log_level == "DEBUG"
    assert model.format_filters == ["docx", "txt"]
    assert model.output_same_dir is False
    assert model.generate_report is False


# ===== A-06: UI 回调异常隔离 =====


def test_process_pending_events_isolates_callback_exceptions():
    """A-06: 单个回调异常不阻断后续事件排空。"""
    from docmask.ui.controller import TaskController
    from docmask.ui.state import AppState

    state = AppState()
    ctrl = TaskController(state, tk_root=None)

    results = []
    def good_callback(value):
        results.append(value)

    ctrl._event_queue.put((lambda: (_ for _ in ()).throw(ValueError("boom")), ()))
    ctrl._event_queue.put((good_callback, ("after_error",)))

    processed = ctrl.process_pending_events()

    assert processed == 2  # 两个回调都被处理
    assert results == ["after_error"]  # 第二个回调正常执行


# ===== A-07: CLI 中断返回码 =====


def test_cli_mask_interrupt_returns_130(tmp_path, monkeypatch):
    """A-07: mask 中断时返回 130，摘要区分中断状态。"""
    import docmask.cli as cli_module

    source = tmp_path / "a.txt"
    source.write_text("content", encoding="utf-8")
    codebook = tmp_path / "cb.txt"
    codebook.write_text("content==>X\n", encoding="utf-8")

    # 让第一个文件处理时抛出 KeyboardInterrupt
    def fake_mask(self, filepath, masker, output_path=None, **kw):
        raise KeyboardInterrupt
    monkeypatch.setattr("docmask.handlers.txt_handler.TxtHandler.mask", fake_mask)

    args = SimpleNamespace(
        input=str(source), codebook=str(codebook), output=None,
        format=None, report=False, allow_empty=False,
    )
    rc = cli_module.cmd_mask(args)
    assert rc == 130


def test_cli_restore_interrupt_returns_130(tmp_path, monkeypatch):
    """A-07: restore 中断时返回 130。"""
    import docmask.cli as cli_module

    source = tmp_path / "a.txt"
    source.write_text("content", encoding="utf-8")
    codebook = tmp_path / "cb.txt"
    codebook.write_text("content==>X\n", encoding="utf-8")

    def fake_restore(self, filepath, restorer, output_path=None, **kw):
        raise KeyboardInterrupt
    monkeypatch.setattr("docmask.handlers.txt_handler.TxtHandler.restore", fake_restore)

    args = SimpleNamespace(
        input=str(source), codebook=str(codebook), output=None,
        format=None, allow_empty=False,
    )
    rc = cli_module.cmd_restore(args)
    assert rc == 130
