"""进度回调与任务取消测试"""
import os
import tempfile
import pytest
from docx import Document

from docmask.core.codebook import Codebook
from docmask.core.masker import Masker
from docmask.core.restorer import Restorer
from docmask.handlers.base import CancelToken, TaskCancelledError
from docmask.handlers.txt_handler import TxtHandler
from docmask.handlers.docx_handler import DocxHandler

DATA_DIR = os.path.join(os.path.dirname(__file__), "test_data")


@pytest.fixture
def codebook():
    cb = Codebook(os.path.join(DATA_DIR, "sample_codebook.txt"))
    cb.load()
    return cb


@pytest.fixture
def masker(codebook):
    return Masker(codebook)


@pytest.fixture
def restorer(codebook):
    return Restorer(codebook)


def _cleanup(paths):
    """安全清理文件列表"""
    for p in paths:
        if p and os.path.exists(p):
            try:
                os.remove(p)
            except OSError:
                pass


class TestCancelToken:
    """测试 CancelToken 基础行为"""

    def test_initial_state_not_cancelled(self):
        token = CancelToken()
        assert not token.is_cancelled

    def test_cancel_sets_flag(self):
        token = CancelToken()
        token.cancel()
        assert token.is_cancelled

    def test_cancel_is_idempotent(self):
        token = CancelToken()
        token.cancel()
        token.cancel()
        assert token.is_cancelled

    def test_thread_safe(self):
        """多线程调用 cancel 不出错"""
        import threading

        token = CancelToken()
        threads = [threading.Thread(target=token.cancel) for _ in range(10)]
        for t in threads:
            t.start()
        for t in threads:
            t.join()
        assert token.is_cancelled


class TestTxtHandlerProgress:
    """测试 TXT Handler 进度回调"""

    def test_mask_reports_progress(self, masker):
        """脱敏时报告完整进度序列"""
        handler = TxtHandler()
        with tempfile.NamedTemporaryFile(
            suffix=".txt", mode="w", delete=False, encoding="utf-8"
        ) as f:
            f.write("张三和李明")
            tmp_in = f.name

        progress_records = []
        output_path = None
        try:
            output_path, _, _ = handler.mask(
                tmp_in, masker,
                progress_callback=lambda c, t, m: progress_records.append((c, t, m)),
            )
            # 应报告开始(0/3)和完成(3/3)
            assert progress_records[0] == (0, 3, "正在读取文件...")
            assert progress_records[-1] == (3, 3, "脱敏完成")
            # 所有 total 应为 3
            for _, total, _ in progress_records:
                assert total == 3
            # current 应单调递增
            currents = [c for c, _, _ in progress_records]
            assert currents == sorted(currents)
        finally:
            _cleanup([tmp_in, output_path])

    def test_restore_reports_progress(self, masker, restorer):
        """恢复时报告完整进度序列"""
        handler = TxtHandler()
        with tempfile.NamedTemporaryFile(
            suffix=".txt", mode="w", delete=False, encoding="utf-8"
        ) as f:
            f.write("张三和李明")
            tmp_in = f.name

        progress_records = []
        masked_path = None
        restored_path = None
        try:
            masked_path, _, _ = handler.mask(tmp_in, masker)
            restored_path, _ = handler.restore(
                masked_path, restorer,
                progress_callback=lambda c, t, m: progress_records.append((c, t, m)),
            )
            assert progress_records[0] == (0, 3, "正在读取文件...")
            assert progress_records[-1] == (3, 3, "恢复完成")
        finally:
            _cleanup([tmp_in, masked_path, restored_path])

    def test_mask_without_callback_works(self, masker):
        """不传 progress_callback 时正常工作"""
        handler = TxtHandler()
        with tempfile.NamedTemporaryFile(
            suffix=".txt", mode="w", delete=False, encoding="utf-8"
        ) as f:
            f.write("张三")
            tmp_in = f.name

        output_path = None
        try:
            output_path, count, _ = handler.mask(tmp_in, masker)
            assert count == 1
        finally:
            _cleanup([tmp_in, output_path])


class TestTxtHandlerCancel:
    """测试 TXT Handler 任务取消"""

    def test_cancel_before_mask_raises(self, masker):
        """脱敏前取消应抛出 TaskCancelledError"""
        handler = TxtHandler()
        with tempfile.NamedTemporaryFile(
            suffix=".txt", mode="w", delete=False, encoding="utf-8"
        ) as f:
            f.write("张三")
            tmp_in = f.name

        token = CancelToken()
        token.cancel()

        try:
            with pytest.raises(TaskCancelledError):
                handler.mask(tmp_in, masker, cancel_token=token)
        finally:
            _cleanup([tmp_in])

    def test_cancel_before_restore_raises(self, restorer):
        """恢复前取消应抛出 TaskCancelledError"""
        handler = TxtHandler()
        token = CancelToken()
        token.cancel()

        with pytest.raises(TaskCancelledError):
            handler.restore("/dev/null", restorer, cancel_token=token)


class TestDocxHandlerProgress:
    """测试 DOCX Handler 进度回调"""

    def test_mask_reports_9_steps(self, masker):
        """脱敏应报告 9 步进度（含 OPC 扩展部件检查）"""
        doc = Document()
        doc.core_properties.author = "张三"
        doc.add_paragraph("张三的工作报告")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_in = f.name

        progress_records = []
        handler = DocxHandler()
        output_path = None
        try:
            output_path, _, _ = handler.mask(
                tmp_in, masker,
                progress_callback=lambda c, t, m: progress_records.append((c, t, m)),
            )
            assert all(t == 9 for _, t, _ in progress_records)
            # 应从 0 开始，到 8 结束
            assert progress_records[0][0] == 0
            assert progress_records[-1] == (9, 9, "OPC 扩展部件检查完成")
        finally:
            _cleanup([tmp_in, output_path])

    def test_restore_reports_8_steps(self, masker, restorer):
        """恢复应报告 8 步进度（含 OPC 扩展部件）"""
        doc = Document()
        doc.core_properties.author = "张三"
        doc.add_paragraph("张三的工作报告")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_in = f.name

        handler = DocxHandler()
        progress_records = []
        masked_path = None
        restored_path = None
        try:
            masked_path, _, _ = handler.mask(tmp_in, masker)
            restored_path, _ = handler.restore(
                masked_path, restorer,
                progress_callback=lambda c, t, m: progress_records.append((c, t, m)),
            )
            assert all(t == 8 for _, t, _ in progress_records)
            assert progress_records[0][0] == 1  # restore 从 1 开始
            assert progress_records[-1] == (8, 8, "OPC 扩展部件恢复完成")
        finally:
            _cleanup([tmp_in, masked_path, restored_path])

    def test_mask_without_callback_works(self, masker):
        """不传 progress_callback 时正常工作"""
        doc = Document()
        doc.add_paragraph("张三")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_in = f.name

        handler = DocxHandler()
        output_path = None
        try:
            output_path, count, _ = handler.mask(tmp_in, masker)
            assert count > 0
        finally:
            _cleanup([tmp_in, output_path])


class TestDocxHandlerCancel:
    """测试 DOCX Handler 任务取消"""

    def test_cancel_during_mask_raises(self, masker):
        """在脱敏过程中取消应抛出 TaskCancelledError"""
        doc = Document()
        doc.add_paragraph("张三")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_in = f.name

        token = CancelToken()
        token.cancel()  # 在开始前就取消

        handler = DocxHandler()
        try:
            with pytest.raises(TaskCancelledError):
                handler.mask(tmp_in, masker, cancel_token=token)
        finally:
            _cleanup([tmp_in])

    def test_cancel_during_restore_raises(self, restorer):
        """在恢复过程中取消应抛出 TaskCancelledError"""
        doc = Document()
        doc.add_paragraph("张三")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_in = f.name

        token = CancelToken()
        token.cancel()

        handler = DocxHandler()
        try:
            with pytest.raises(TaskCancelledError):
                handler.restore(tmp_in, restorer, cancel_token=token)
        finally:
            _cleanup([tmp_in])
