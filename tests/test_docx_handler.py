"""DOCX 文件处理器测试"""
import os
import tempfile
import pytest
from docx import Document
from docmask.core.codebook import Codebook
from docmask.core.masker import Masker
from docmask.core.restorer import Restorer
from docmask.handlers.docx_handler import DocxHandler

DATA_DIR = os.path.join(os.path.dirname(__file__), "test_data")


@pytest.fixture
def codebook():
    cb = Codebook(os.path.join(DATA_DIR, "sample_codebook.txt"))
    cb.load()
    return cb


@pytest.fixture
def masker(codebook):
    return Masker(codebook)


@pytest.fixture
def restorer(codebook):
    return Restorer(codebook)


def _create_test_docx():
    """创建测试用 docx"""
    doc = Document()
    doc.core_properties.author = "张三"
    doc.core_properties.title = "某某科技有限公司报告"
    doc.add_paragraph("张三的工作报告")
    return doc


class TestDocxHandler:
    """测试 DOCX 文件处理器"""

    def test_mask_paragraphs(self, masker):
        """正文段落脱敏"""
        doc = _create_test_docx()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_in = f.name

        handler = DocxHandler()
        try:
            output_path, count, coverage = handler.mask(tmp_in, masker)
            assert count > 0
            # 验证脱敏结果
            result_doc = Document(output_path)
            texts = [p.text for p in result_doc.paragraphs]
            combined = "".join(texts)
            assert "李四" in combined
            assert "张三" not in combined
        finally:
            os.remove(tmp_in)
            if os.path.exists(output_path):
                os.remove(output_path)

    def test_mask_preserves_format(self, masker):
        """格式保留：加粗、字体等"""
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("张三")
        run.bold = True
        run.font.size = None  # 使用默认

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_in = f.name

        handler = DocxHandler()
        try:
            output_path, count, _ = handler.mask(tmp_in, masker)
            result_doc = Document(output_path)
            result_run = result_doc.paragraphs[0].runs[0]
            assert result_run.bold is True
            assert result_run.text == "李四"
        finally:
            os.remove(tmp_in)
            if os.path.exists(output_path):
                os.remove(output_path)

    def test_mask_metadata(self, masker):
        """元数据脱敏"""
        doc = _create_test_docx()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_in = f.name

        handler = DocxHandler()
        try:
            output_path, count, _ = handler.mask(tmp_in, masker)
            result_doc = Document(output_path)
            assert result_doc.core_properties.author == "李四"
        finally:
            os.remove(tmp_in)
            if os.path.exists(output_path):
                os.remove(output_path)

    def test_restore_roundtrip(self, masker, restorer):
        """脱敏后恢复 roundtrip"""
        doc = _create_test_docx()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_in = f.name

        handler = DocxHandler()
        try:
            # 脱敏
            masked_path, _, _ = handler.mask(tmp_in, masker)
            # 恢复
            restored_path, count = handler.restore(masked_path, restorer)
            result_doc = Document(restored_path)
            texts = [p.text for p in result_doc.paragraphs]
            combined = "".join(texts)
            # 精确规则应恢复
            assert "张三" in combined
            assert result_doc.core_properties.author == "张三"
        finally:
            for p in [tmp_in, masked_path, restored_path]:
                if os.path.exists(p):
                    os.remove(p)

    def test_mask_tables(self, masker):
        """表格文本脱敏"""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "姓名"
        table.cell(0, 1).text = "张三"

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_in = f.name

        handler = DocxHandler()
        try:
            output_path, count, _ = handler.mask(tmp_in, masker)
            result_doc = Document(output_path)
            assert result_doc.tables[0].cell(0, 1).text == "李四"
        finally:
            os.remove(tmp_in)
            if os.path.exists(output_path):
                os.remove(output_path)

    def test_mask_headers(self, masker):
        """页眉脱敏"""
        doc = Document()
        section = doc.sections[0]
        header = section.header
        header.paragraphs[0].text = "某某科技有限公司"

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_in = f.name

        handler = DocxHandler()
        try:
            output_path, count, _ = handler.mask(tmp_in, masker)
            result_doc = Document(output_path)
            header_text = result_doc.sections[0].header.paragraphs[0].text
            assert "甲乙科技有限公司" in header_text
        finally:
            os.remove(tmp_in)
            if os.path.exists(output_path):
                os.remove(output_path)