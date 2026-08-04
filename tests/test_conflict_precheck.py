"""测试文档级脱敏词冲突预检功能"""
import os
import tempfile
import pytest

from docmask.core.codebook import Codebook
from docmask.core.masker import Masker, MaskConflictError
from docmask.handlers.txt_handler import TxtHandler
from docmask.handlers.docx_handler import DocxHandler


class TestPrecheckConflict:
    """测试 Masker.precheck_conflict 方法"""

    def _make_codebook(self, rules: list[str], tmp_path) -> Codebook:
        """创建临时密码本"""
        cb_path = os.path.join(tmp_path, "codebook.txt")
        with open(cb_path, "w", encoding="utf-8") as f:
            f.write("\n".join(rules))
        cb = Codebook(cb_path)
        cb.load()
        return cb

    def test_no_conflict(self, tmp_path):
        """脱敏词不在文本中 → 无冲突"""
        cb = self._make_codebook(["张三==>李四", "北京==>上海"], tmp_path)
        masker = Masker(cb)
        conflicts = masker.precheck_conflict("张三住在北京")
        assert conflicts == []

    def test_exact_conflict(self, tmp_path):
        """脱敏词已存在于文本中 → 检测到冲突"""
        cb = self._make_codebook(["6==>%"], tmp_path)
        masker = Masker(cb)
        # 文本中包含 %（脱敏词），应检测到冲突
        conflicts = masker.precheck_conflict("占比80%")
        assert len(conflicts) == 1
        assert conflicts[0][0] == "%"
        assert conflicts[0][1] == 1

    def test_multiple_conflicts(self, tmp_path):
        """多个脱敏词同时冲突"""
        cb = self._make_codebook(
            ["6==>%", "8==>$", "安全==>保护"], tmp_path
        )
        masker = Masker(cb)
        # 文本中包含 % 和 $（脱敏词），但不含 保护
        conflicts = masker.precheck_conflict("费用%和$以及安全")
        conflict_words = [c[0] for c in conflicts]
        assert "%" in conflict_words
        assert "$" in conflict_words
        assert "保护" not in conflict_words  # "保护"不在文本中

    def test_conflict_count(self, tmp_path):
        """冲突次数统计正确"""
        cb = self._make_codebook(["a==>x"], tmp_path)
        masker = Masker(cb)
        conflicts = masker.precheck_conflict("xyz xyz xyz")
        assert conflicts[0][1] == 3  # "x" 出现3次

    def test_regex_replacement_conflict(self, tmp_path):
        """正则规则的脱敏词冲突也能检测"""
        cb = self._make_codebook(
            ["regex:\\d{4}年==>脱敏年份"], tmp_path
        )
        masker = Masker(cb)
        # 文本中包含"脱敏年份"（正则规则的脱敏词）
        conflicts = masker.precheck_conflict("这是脱敏年份的记录")
        assert len(conflicts) == 1
        assert conflicts[0][0] == "脱敏年份"

    def test_duplicate_replacement_checked_once(self, tmp_path):
        """相同脱敏词只检查一次（精确规则和正则规则共用同一脱敏词）"""
        cb = self._make_codebook(
            ["甲==>XX", "regex:\\d+==>XX"], tmp_path
        )
        masker = Masker(cb)
        conflicts = masker.precheck_conflict("文本中有XX")
        assert len(conflicts) == 1
        assert conflicts[0][0] == "XX"

    def test_empty_codebook_no_conflict(self, tmp_path):
        """空文本不产生冲突"""
        cb = self._make_codebook(["张三==>李四"], tmp_path)
        masker = Masker(cb)
        conflicts = masker.precheck_conflict("")
        assert conflicts == []


class TestTxtHandlerConflict:
    """测试 TxtHandler 的冲突预检"""

    def _setup(self, tmp_path, codebook_rules, text_content):
        cb_path = os.path.join(tmp_path, "codebook.txt")
        with open(cb_path, "w", encoding="utf-8") as f:
            f.write("\n".join(codebook_rules))
        cb = Codebook(cb_path)
        cb.load()

        txt_path = os.path.join(tmp_path, "input.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(text_content)

        return cb, txt_path

    def test_txt_mask_blocked_on_conflict(self, tmp_path):
        """TXT 脱敏时检测到冲突 → 抛出 MaskConflictError"""
        cb, txt_path = self._setup(
            tmp_path, ["6==>%"], "2015年占比80%"
        )
        masker = Masker(cb)
        handler = TxtHandler()

        with pytest.raises(MaskConflictError) as exc_info:
            handler.mask(txt_path, masker)

        assert "%" in str(exc_info.value)
        assert "已存在脱敏词" in str(exc_info.value)

    def test_txt_mask_no_conflict_passes(self, tmp_path):
        """TXT 脱敏时无冲突 → 正常执行"""
        cb, txt_path = self._setup(
            tmp_path, ["张三==>李四"], "张三去北京了"
        )
        masker = Masker(cb)
        handler = TxtHandler()

        output_path, count, _ = handler.mask(txt_path, masker)
        assert count == 1
        assert os.path.exists(output_path)


class TestDocxHandlerConflict:
    """测试 DocxHandler 的冲突预检"""

    def test_docx_mask_blocked_on_conflict(self, tmp_path):
        """DOCX 脱敏时检测到冲突 → 抛出 MaskConflictError"""
        from docx import Document

        # 创建密码本：6==>%，脱敏词 % 会在原文中出现
        cb_path = os.path.join(tmp_path, "codebook.txt")
        with open(cb_path, "w", encoding="utf-8") as f:
            f.write("6==>%\n8==>$")
        cb = Codebook(cb_path)
        cb.load()

        # 创建包含 % 的 docx 文档
        docx_path = os.path.join(tmp_path, "input.docx")
        doc = Document()
        doc.add_paragraph("2015年占比80%")
        doc.save(docx_path)

        masker = Masker(cb)
        handler = DocxHandler()

        with pytest.raises(MaskConflictError) as exc_info:
            handler.mask(docx_path, masker)

        assert "%" in str(exc_info.value)

    def test_docx_mask_no_conflict_passes(self, tmp_path):
        """DOCX 脱敏时无冲突 → 正常执行"""
        from docx import Document

        cb_path = os.path.join(tmp_path, "codebook.txt")
        with open(cb_path, "w", encoding="utf-8") as f:
            f.write("张三==>李四")
        cb = Codebook(cb_path)
        cb.load()

        docx_path = os.path.join(tmp_path, "input.docx")
        doc = Document()
        doc.add_paragraph("张三去北京了")
        doc.save(docx_path)

        masker = Masker(cb)
        handler = DocxHandler()

        output_path, count, _ = handler.mask(docx_path, masker)
        assert count == 1
        assert os.path.exists(output_path)
