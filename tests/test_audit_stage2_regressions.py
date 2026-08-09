"""Stage 2 审计修复回归测试。

覆盖：A-08 事务性保存 / A-09 批量历史写入 / A-10 可取消 DOC 转换 /
      A-11 文件大小预算 / A-12 不可变 TaskContext / A-17 恢复残留校验
"""
from __future__ import annotations

import json
import os
import shutil
import subprocess
import threading
import time
import zipfile
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

import pytest
from docx import Document

from docmask.core.codebook import Codebook, CodebookRule
from docmask.core.masker import Masker, _MAX_CANDIDATES, CandidateBudgetExceededError
from docmask.core.restorer import Restorer
from docmask.handlers.base import CancelToken, TaskCancelledError
from docmask.handlers.doc_handler import DocHandler
from docmask.handlers.docx_handler import DocxHandler
from docmask.handlers.txt_handler import TxtHandler, TxtFileTooLargeError, _MAX_TXT_FILE_SIZE
from docmask.services.codebook_library import CodebookLibrary
from docmask.services.history_store import HistoryStore, HistoryEntry
from docmask.ui.controller import TaskController
from docmask.ui.state import AppState, FileStatus, Mode, TaskContext, create_file_item


def make_codebook(tmp_path: Path, rules: str) -> Codebook:
    path = tmp_path / "codebook.txt"
    path.write_text(rules, encoding="utf-8")
    cb = Codebook(str(path))
    cb.load()
    return cb


def make_entry(timestamp: str, filename: str = "报告.docx", replacements: int = 10):
    return HistoryEntry(
        timestamp=timestamp,
        mode="mask",
        input_path=f"/test/{filename}",
        input_filename=filename,
        output_path=f"/test/{filename}_desensitized",
        codebook_name="默认密码本",
        codebook_version="v-20260808_150000",
        exact_rule_count=15,
        regex_rule_count=2,
        replacements=replacements,
        status="done",
    )


# ===== A-08: 事务性密码本保存 =====


class TestTransactionalSave:
    """A-08: commit marker 事务性保存测试。"""

    def test_normal_save_no_commit_marker(self, tmp_path):
        library = CodebookLibrary(base_dir=tmp_path / "codebooks")
        cb = make_codebook(tmp_path, "张三==>李四\n")
        meta = library.create("测试")
        library.save(meta.id, cb)
        cb_dir = library._codebook_dir(meta.id)
        assert not (cb_dir / ".commit").exists()

    def test_recover_from_interrupted_save(self, tmp_path):
        """模拟保存中断（commit marker 存在），验证启动时恢复。"""
        library = CodebookLibrary(base_dir=tmp_path / "codebooks")
        cb = make_codebook(tmp_path, "张三==>李四\n北京市==>⟦DM-ADDR-01⟧\n")
        meta = library.create("测试")
        v1 = library.save(meta.id, cb)

        # 模拟第二次保存被中断：写入版本文件和 commit marker，但不更新 current.txt
        new_content = "王五==>赵六\n"
        version_id = "v-test_interrupted"
        old_meta = library._read_meta(meta.id)
        new_meta = dict(old_meta)
        new_meta["current_version"] = version_id
        new_meta["versions"].append({
            "version_id": version_id,
            "created_at": "2026-01-01T00:00:00",
            "exact_rule_count": 1,
            "regex_rule_count": 0,
            "change_summary": "+1 精确规则",
        })

        version_path = library._versions_dir(meta.id) / f"{version_id}.txt"
        version_path.write_text(new_content, encoding="utf-8")

        commit_marker = library._codebook_dir(meta.id) / ".commit"
        commit_marker.write_text(
            json.dumps(new_meta, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        # current.txt 仍是旧内容（模拟中断在写 current.txt 之前）
        old_current = library._current_path(meta.id).read_text(encoding="utf-8")
        assert "张三" in old_current

        # 创建新实例，触发恢复
        library2 = CodebookLibrary(base_dir=library._base_dir)

        # 验证恢复：current.txt 应被更新为新内容
        loaded = library2.load(meta.id)
        assert loaded.exact_rule_count == 1
        assert "王五" in loaded.forward_map

        # commit marker 应被删除
        assert not commit_marker.exists()

    def test_recover_with_missing_version_file(self, tmp_path):
        """commit marker 存在但版本文件不存在时，旧状态保持完好。"""
        library = CodebookLibrary(base_dir=tmp_path / "codebooks")
        cb = make_codebook(tmp_path, "张三==>李四\n")
        meta = library.create("测试")
        library.save(meta.id, cb)

        # 写入 commit marker 但不写版本文件
        commit_marker = library._codebook_dir(meta.id) / ".commit"
        commit_marker.write_text(
            json.dumps({"current_version": "v-nonexistent"}, ensure_ascii=False),
            encoding="utf-8",
        )

        old_current = library._current_path(meta.id).read_text(encoding="utf-8")

        library2 = CodebookLibrary(base_dir=library._base_dir)

        # 旧状态应保持完好
        loaded = library2.load(meta.id)
        assert loaded.exact_rule_count == 1
        assert "张三" in loaded.forward_map
        assert not commit_marker.exists()


# ===== A-09: 批量历史写入 =====


class TestAppendMany:
    """A-09: append_many 批量写入测试。"""

    def test_append_many_count_correct(self, tmp_path):
        store = HistoryStore(path=tmp_path / "history.jsonl")
        entries = [make_entry(f"2026-08-08T10:{i:04d}:00") for i in range(100)]
        store.append_many(entries)
        assert store.count() == 100

    def test_append_many_large_batch(self, tmp_path):
        store = HistoryStore(path=tmp_path / "history.jsonl")
        entries = [make_entry(f"2026-08-08T{i:05d}:00") for i in range(1000)]
        store.append_many(entries)
        assert store.count() == 1000

    def test_append_many_trims_to_max(self, tmp_path):
        store = HistoryStore(path=tmp_path / "history.jsonl")
        for i in range(500):
            store.append(make_entry(f"2026-08-08T{i:04d}:00"))
        entries = [make_entry(f"2026-08-09T{i:05d}:00") for i in range(800)]
        store.append_many(entries)
        assert store.count() == HistoryStore.MAX_ENTRIES

    def test_append_many_empty(self, tmp_path):
        store = HistoryStore(path=tmp_path / "history.jsonl")
        store.append_many([])
        assert store.count() == 0

    def test_append_many_preserves_order(self, tmp_path):
        store = HistoryStore(path=tmp_path / "history.jsonl")
        entries = [
            make_entry(f"2026-08-08T10:0{i}:00", filename=f"file{i}.txt")
            for i in range(5)
        ]
        store.append_many(entries)
        result = store.query(limit=5)
        assert result[0].input_filename == "file4.txt"
        assert result[4].input_filename == "file0.txt"

    def test_append_many_after_existing(self, tmp_path):
        store = HistoryStore(path=tmp_path / "history.jsonl")
        store.append(make_entry("2026-08-08T10:00:00", filename="old.txt"))
        entries = [make_entry(f"2026-08-08T11:0{i}:00", filename=f"new{i}.txt") for i in range(3)]
        store.append_many(entries)
        assert store.count() == 4
        result = store.query(limit=4)
        assert result[0].input_filename == "new2.txt"
        assert result[3].input_filename == "old.txt"


# ===== A-10: 可取消 DOC 转换 =====


def test_libreoffice_convert_cancellable(tmp_path, monkeypatch):
    """A-10: 取消令牌触发后 LibreOffice 子进程被终止。"""
    legacy_doc = tmp_path / "source.doc"
    legacy_doc.write_bytes(b"placeholder")
    output_path = tmp_path / "source.docx"
    handler = DocHandler()

    monkeypatch.setattr(handler, "_find_libreoffice_command", lambda: "soffice")

    terminate_called = {"value": False}

    class SlowPopen:
        def __init__(self, args, **_kwargs):
            self._args = args
            self.returncode = None
            self._terminated = False

        def communicate(self, timeout=None):
            if self._terminated:
                return ("", "")
            raise subprocess.TimeoutExpired(cmd=self._args, timeout=timeout)

        def poll(self):
            return self.returncode

        def terminate(self):
            terminate_called["value"] = True
            self._terminated = True
            self.returncode = -15

        def kill(self):
            self._terminated = True
            self.returncode = -9

    monkeypatch.setattr("docmask.handlers.doc_handler.subprocess.Popen", SlowPopen)

    cancel_token = CancelToken()
    cancel_token.cancel()

    with pytest.raises(TaskCancelledError):
        handler._try_libreoffice_convert(
            str(legacy_doc), output_path, cancel_token
        )

    assert terminate_called["value"]


# ===== A-11: 文件大小预算 + 有界候选 =====


def test_candidate_budget_exceeded(tmp_path):
    """A-11: 候选匹配数量超过预算时抛出异常。"""
    codebook = make_codebook(tmp_path, "a==>X\n")
    masker = Masker(codebook)
    text = "a " * (_MAX_CANDIDATES + 10)

    with pytest.raises(CandidateBudgetExceededError):
        masker.mask_text(text)


def test_candidate_at_boundary_ok(tmp_path):
    """A-11: 候选数量恰好等于上限时正常处理。"""
    codebook = make_codebook(tmp_path, "a==>X\n")
    masker = Masker(codebook)
    text = "a " * _MAX_CANDIDATES
    _, count, _ = masker.mask_text(text)
    assert count == _MAX_CANDIDATES


def test_txt_file_too_large_rejected(tmp_path):
    """A-11: 超大 TXT 文件拒绝处理。"""
    large_file = tmp_path / "large.txt"
    large_file.write_text("content", encoding="utf-8")

    with patch("os.path.getsize", return_value=_MAX_TXT_FILE_SIZE + 1):
        with pytest.raises(TxtFileTooLargeError):
            TxtHandler().read(str(large_file))


def test_txt_file_at_size_boundary_ok(tmp_path):
    """A-11: 文件大小恰好等于上限时正常处理。"""
    large_file = tmp_path / "ok.txt"
    large_file.write_text("content", encoding="utf-8")

    with patch("os.path.getsize", return_value=_MAX_TXT_FILE_SIZE):
        content = TxtHandler().read(str(large_file))
        assert content == "content"


# ===== A-12: 不可变 TaskContext =====


class ImmediateTkRoot:
    def __init__(self):
        self.owner_thread = threading.get_ident()
        self.callbacks = []

    def after(self, _delay, callback):
        assert threading.get_ident() == self.owner_thread
        self.callbacks.append(callback)

    def flush(self):
        while self.callbacks:
            self.callbacks.pop(0)()


def _make_state(tmp_path: Path, output_same_dir: bool = True) -> AppState:
    codebook_path = tmp_path / "codebook.txt"
    codebook_path.write_text("张三==>李四\n", encoding="utf-8")
    codebook = Codebook(str(codebook_path))
    codebook.load()

    input_path = tmp_path / "input.txt"
    input_path.write_text("张三的文件", encoding="utf-8")

    state = AppState()
    state.codebook.codebook = codebook
    state.codebook.valid = True
    state.codebook.path = str(codebook_path)
    state.files = []
    state.history_enabled = False
    state.output_same_dir = output_same_dir
    state.output_dir = str(tmp_path / "output") if not output_same_dir else None
    if state.output_dir:
        Path(state.output_dir).mkdir()
    return state


def _run(controller: TaskController):
    progress = []
    completed = threading.Event()
    started = controller.execute(
        on_file_start=lambda index, item: None,
        on_file_done=lambda index, item: None,
        on_progress=lambda current, total, message: progress.append(
            (current, total, message)
        ),
        on_complete=lambda results: completed.set(),
    )
    assert started is True
    controller._thread.join(timeout=5)
    assert not controller._thread.is_alive()
    controller.tk_root.flush()
    assert completed.is_set()
    return progress


def test_task_context_is_frozen(tmp_path):
    """A-12: TaskContext 是不可变 dataclass。"""
    codebook = make_codebook(tmp_path, "张三==>李四\n")
    ctx = TaskContext(
        mode=Mode.MASK,
        codebook=codebook,
        codebook_name="原始密码本",
        codebook_version="v-1",
        exact_count=1,
        regex_count=0,
        output_same_dir=True,
        output_dir=None,
        generate_report=True,
        history_enabled=True,
    )
    with pytest.raises(AttributeError):
        ctx.codebook_name = "新密码本"
    assert ctx.codebook_name == "原始密码本"


def test_task_context_snapshot_not_affected_by_state_change(tmp_path):
    """A-12: 任务完成后 TaskContext 快照不受后续状态修改影响。"""
    state = _make_state(tmp_path)
    state.codebook.library_name = "原始密码本"
    input_path = tmp_path / "input.txt"
    state.files = [create_file_item(str(input_path))]
    controller = TaskController(state, ImmediateTkRoot())

    _run(controller)

    assert controller._task_context is not None
    original_name = controller._task_context.codebook_name

    # 修改 state，TaskContext 不应改变
    state.codebook.library_name = "修改后的密码本"
    assert controller._task_context.codebook_name == original_name


def test_history_uses_task_context_not_current_state(tmp_path):
    """A-12: 历史记录使用任务开始时的 TaskContext，不受后续状态修改影响。"""
    state = _make_state(tmp_path)
    state.history_enabled = True
    state.codebook.library_name = "原始密码本"
    input_path = tmp_path / "input.txt"
    state.files = [create_file_item(str(input_path))]
    controller = TaskController(state, ImmediateTkRoot())

    _run(controller)

    # 第一次 flush 触发 _schedule_record_history
    # 修改 state（模拟任务后用户切换密码本）
    state.codebook.library_name = "新密码本"

    # 第二次 flush 触发 record_history（使用 TaskContext 快照）
    controller.tk_root.flush()

    # 等待后台历史写入线程完成
    time.sleep(0.2)

    history = controller.query_history(limit=1)
    assert len(history) == 1
    # 应使用任务开始时的密码本名，不是修改后的
    assert history[0].codebook_name == "原始密码本"


# ===== A-17: DOCX 恢复后残留校验 =====


def test_restore_residual_check_rejects_masked_word(tmp_path):
    """A-17: 恢复后仍有脱敏词残留时拒绝提交。"""
    source = tmp_path / "masked.docx"
    doc = Document()
    doc.add_paragraph("李四的文件")
    doc.save(source)

    codebook = make_codebook(tmp_path, "张三==>李四\n")
    restorer = Restorer(codebook)

    # 模拟恢复失败：restore_text 不替换任何内容
    restorer.restore_text = lambda text: (text, 0)

    handler = DocxHandler()
    with pytest.raises(RuntimeError, match="恢复后残留校验失败"):
        handler.restore(str(source), restorer)

    # 不应产生输出文件
    assert not list(tmp_path.glob("*_restored*"))


def test_restore_normal_passes_residual_check(tmp_path):
    """A-17: 正常恢复通过残留校验。"""
    source = tmp_path / "masked.docx"
    doc = Document()
    doc.add_paragraph("李四的文件")
    doc.save(source)

    codebook = make_codebook(tmp_path, "张三==>李四\n")
    restorer = Restorer(codebook)
    handler = DocxHandler()

    output_path, count = handler.restore(str(source), restorer)
    assert count == 1
    assert Path(output_path).exists()

    # 验证恢复结果
    result = Document(output_path)
    text = "\n".join(p.text for p in result.paragraphs)
    assert "张三" in text
    assert "李四" not in text


def test_restore_no_masked_words_passes_check(tmp_path):
    """A-17: 文档中无脱敏词时残留校验通过。"""
    source = tmp_path / "plain.docx"
    doc = Document()
    doc.add_paragraph("普通文本，无脱敏词")
    doc.save(source)

    codebook = make_codebook(tmp_path, "张三==>李四\n")
    restorer = Restorer(codebook)
    handler = DocxHandler()

    output_path, count = handler.restore(str(source), restorer)
    assert count == 0
    assert Path(output_path).exists()
