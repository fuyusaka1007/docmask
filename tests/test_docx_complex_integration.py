"""复杂 DOCX 文本节点、超链接和 note part 集成测试。"""
from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree

from docmask.core.codebook import Codebook
from docmask.core.masker import Masker
from docmask.core.restorer import Restorer
from docmask.handlers.docx_handler import DocxHandler, NSMAP
from tests.word_fixture_factory import create_complex_word_fixture


def _engines(tmp_path: Path):
    codebook_path = tmp_path / "codebook.txt"
    codebook_path.write_text("张三==>匿名客户001\n", encoding="utf-8")
    codebook = Codebook(str(codebook_path))
    codebook.load()
    return Masker(codebook), Restorer(codebook)


def _part_xml(path: str | Path, member: str) -> bytes:
    with ZipFile(path) as package:
        return package.read(member)


def _structural_counts(path: str | Path) -> dict[str, int]:
    root = etree.fromstring(_part_xml(path, "word/document.xml"))
    return {
        "tabs": len(root.findall(".//w:tab", NSMAP)),
        "breaks": len(root.findall(".//w:br", NSMAP)),
        "drawings": len(root.findall(".//w:drawing", NSMAP)),
        "field_chars": len(root.findall(".//w:fldChar", NSMAP)),
        "field_instructions": len(root.findall(".//w:instrText", NSMAP)),
    }


def test_complex_docx_mask_restore_preserves_structure_and_updates_all_text(tmp_path):
    source = create_complex_word_fixture(tmp_path / "complex_source.docx")
    masker, restorer = _engines(tmp_path)
    handler = DocxHandler()

    source_structure = _structural_counts(source)
    masked, count, _coverage = handler.mask(
        str(source), masker, output_path=str(tmp_path / "complex_masked.docx"),
    )
    restored, restored_count = handler.restore(
        masked, restorer, output_path=str(tmp_path / "complex_restored.docx"),
    )

    assert count >= 7
    assert restored_count == count
    assert _structural_counts(masked) == source_structure
    assert _structural_counts(restored) == source_structure

    masked_document = Document(masked)
    link_texts = [
        hyperlink.text
        for paragraph in masked_document.paragraphs
        for hyperlink in paragraph.hyperlinks
    ]
    assert "匿名客户001" in link_texts
    assert "张三" not in link_texts

    restored_document = Document(restored)
    restored_links = [
        hyperlink.text
        for paragraph in restored_document.paragraphs
        for hyperlink in paragraph.hyperlinks
    ]
    assert "张三" in restored_links

    for member in ("word/footnotes.xml", "word/endnotes.xml"):
        masked_xml = _part_xml(masked, member).decode("utf-8")
        restored_xml = _part_xml(restored, member).decode("utf-8")
        assert "匿名客户001" in masked_xml
        assert "张三" not in masked_xml
        assert "张三" in restored_xml


def test_length_changing_cross_run_replacement_keeps_run_properties(tmp_path):
    source = create_complex_word_fixture(tmp_path / "format_source.docx")
    masker, _restorer = _engines(tmp_path)
    output, _count, _coverage = DocxHandler().mask(
        str(source), masker, output_path=str(tmp_path / "format_masked.docx"),
    )

    document = Document(output)
    paragraph = next(p for p in document.paragraphs if p.text.startswith("普通正文"))
    assert "匿名客户001" in paragraph.text
    assert paragraph.runs[1].bold is True
    assert paragraph.runs[2].italic is True
    assert "\t" in paragraph.text
    assert "\n" in paragraph.text

