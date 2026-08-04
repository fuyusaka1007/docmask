"""macOS LibreOffice 旧 DOC 实际转换链路集成测试

依赖条件：
- LibreOffice 已安装（/Applications/LibreOffice.app/Contents/MacOS/soffice）
- 测试文档测试文档.doc 位于项目根目录

使用 pytest 标记 skip_if_no_libreoffice 在 LO 不可用时自动跳过。
"""
from __future__ import annotations

import shutil
import tempfile
import zipfile
from pathlib import Path

import pytest

from docmask.core.codebook import Codebook
from docmask.core.masker import Masker, MaskConflictError
from docmask.core.restorer import Restorer
from docmask.handlers.doc_handler import DocHandler
from docmask.handlers.docx_handler import DocxHandler


PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEST_DOC = PROJECT_ROOT / "测试文档.doc"


def _libreoffice_available() -> bool:
    handler = DocHandler()
    return handler._find_libreoffice_command() is not None


skip_if_no_libreoffice = pytest.mark.skipif(
    not _libreoffice_available(),
    reason="LibreOffice 未安装，跳过 .doc 实际转换测试",
)

skip_if_no_test_doc = pytest.mark.skipif(
    not TEST_DOC.is_file(),
    reason=f"测试文档未找到：{TEST_DOC}",
)


# ======================== 辅助工具 ========================

def make_codebook(tmp_path: Path, rules: str) -> tuple[Masker, Restorer]:
    """根据规则文本创建 Masker 和 Restorer。"""
    path = tmp_path / "codebook.txt"
    path.write_text(rules, encoding="utf-8")
    cb = Codebook(str(path))
    cb.load()
    return Masker(cb), Restorer(cb)


def _runs_text(doc) -> str:
    """提取文档所有段落文本（包含表格）。"""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


def _docx_contains(docx_path: Path, text: str) -> bool:
    """检查 DOCX 文件 ZIP 包内是否包含指定文本。"""
    with zipfile.ZipFile(docx_path) as archive:
        for name in archive.namelist():
            if name.endswith(".xml") or name.endswith(".rels"):
                try:
                    content = archive.read(name).decode("utf-8")
                    if text in content:
                        return True
                except UnicodeDecodeError:
                    pass
    return False


# ======================== LO 转换测试 ========================

@skip_if_no_libreoffice
@skip_if_no_test_doc
def test_libreoffice_converts_doc_to_valid_docx(tmp_path):
    """LibreOffice 能将 测试文档.doc 转换为有效的 DOCX。"""
    handler = DocHandler()
    # 注意：在 with 块内断言，否则 TemporaryDirectory 退出后文件已删除
    with tempfile.TemporaryDirectory(prefix="docmask_lo_test_") as temp_dir:
        converted = handler._convert_to_docx(str(TEST_DOC), temp_dir)
        assert handler._is_valid_docx(converted)
        from docx import Document
        doc = Document(converted)
        assert len(doc.paragraphs) == 17
        assert "投资并购" in _runs_text(doc)


@skip_if_no_libreoffice
@skip_if_no_test_doc
def test_try_libreoffice_convert_returns_correct_path(tmp_path):
    """_try_libreoffice_convert 返回 LO 实际生成的文件路径。

    注意：LO 生成的文件名基于输入 stem（从 _convert_to_docx 调用时二者一致）。
    直接调用 _try_libreoffice_convert 时输出路径的父目录被作为 --outdir，
    实际文件名由 LO 根据输入 stem 决定。
    """
    handler = DocHandler()
    output_path = tmp_path / "测试文档.docx"

    result = handler._try_libreoffice_convert(str(TEST_DOC), output_path)

    # LO 生成的路径 = outdir / input_stem.docx
    assert result == output_path
    assert result.is_file()
    assert result.stat().st_size > 0
    assert zipfile.is_zipfile(result)


@skip_if_no_libreoffice
@skip_if_no_test_doc
def test_convert_to_docx_returns_lo_output(tmp_path):
    """_convert_to_docx 在 macOS 上选择 LibreOffice 并返回真实输出。"""
    handler = DocHandler()
    # 在 macOS 上 pywin32 不可用，应 fallback 到 LO

    converted = handler._convert_to_docx(str(TEST_DOC), str(tmp_path))

    assert Path(converted).is_file()
    assert handler._is_valid_docx(converted)
    assert Path(converted).parent == tmp_path


# ======================== DOC mask → restore 完整链路 ========================

@skip_if_no_libreoffice
@skip_if_no_test_doc
def test_doc_mask_uses_real_libreoffice_and_writes_valid_docx(tmp_path):
    """使用真实 LibreOffice 脱敏 .doc，输出合法 .docx。"""
    masker, _restorer = make_codebook(
        tmp_path,
        "假药集团==>某集团公司\n"
        "李加薪==>张某某\n",
    )
    handler = DocHandler()
    output = tmp_path / "masked.doc"

    masked_path, count, coverage = handler.mask(
        str(TEST_DOC), masker, output_path=str(output),
    )

    assert Path(masked_path).suffix == ".docx"
    assert handler._is_valid_docx(masked_path)
    # 应检测到多次"假药集团"和"李加薪"
    assert count > 0

    # 输出文档中不包含原始敏感内容
    assert not _docx_contains(Path(masked_path), "假药集团")
    assert not _docx_contains(Path(masked_path), "李加薪")
    # 应包含脱敏词
    assert _docx_contains(Path(masked_path), "某集团公司")
    assert _docx_contains(Path(masked_path), "张某某")


@skip_if_no_libreoffice
@skip_if_no_test_doc
def test_doc_mask_restore_round_trip_via_libreoffice(tmp_path):
    """使用真实 LibreOffice 完成 .doc → mask → restore 往返。

    脱敏后得到 .docx，恢复由 DocxHandler 直接处理（无需二次 LO 转换）。
    """
    masker, restorer = make_codebook(
        tmp_path,
        "假药集团==>某集团公司\n"
        "李加薪==>张某某\n",
    )
    doc_handler = DocHandler()
    docx_handler = DocxHandler()
    masked_output = tmp_path / "masked.docx"
    restored_output = tmp_path / "restored.docx"

    # 第一步：脱敏 .doc — 通过 LO 转换 + DocxHandler
    masked_path, mask_count, _coverage = doc_handler.mask(
        str(TEST_DOC), masker, output_path=str(masked_output),
    )
    assert mask_count > 0
    assert not _docx_contains(Path(masked_path), "假药集团")
    assert not _docx_contains(Path(masked_path), "李加薪")

    # 第二步：恢复 — 脱敏后已是 .docx，直接用 DocxHandler
    restored_path, restore_count = docx_handler.restore(
        masked_path, restorer, output_path=str(restored_output),
    )

    assert Path(restored_path).suffix == ".docx"
    assert _docx_contains(Path(restored_path), "假药集团")
    assert _docx_contains(Path(restored_path), "李加薪")
    assert restore_count == mask_count
    # 不应包含脱敏词
    assert not _docx_contains(Path(restored_path), "某集团公司")
    assert not _docx_contains(Path(restored_path), "张某某")


@skip_if_no_libreoffice
@skip_if_no_test_doc
def test_doc_mask_preserves_document_structure(tmp_path):
    """脱敏不破坏文档结构（段落数、非文本内容）。"""
    masker, _restorer = make_codebook(
        tmp_path,
        "假药集团==>MASKED_GROUP\n"
        "李加薪==>MASKED_NAME\n",
    )
    handler = DocHandler()

    masked_path, count, _coverage = handler.mask(
        str(TEST_DOC), masker, output_path=str(tmp_path / "out.doc"),
    )

    from docx import Document
    masked_doc = Document(masked_path)

    # 段落数应保持不变
    assert len(masked_doc.paragraphs) == 17

    # 原始文本结构验证：首段标题不变
    assert masked_doc.paragraphs[0].text == "第X节  投资并购"
    assert "自2022年以来" in masked_doc.paragraphs[2].text
    assert "从零起步" in masked_doc.paragraphs[2].text

    assert count > 0
    assert "MASKED_GROUP" in _runs_text(masked_doc)
    assert "MASKED_NAME" in _runs_text(masked_doc)


@skip_if_no_libreoffice
@skip_if_no_test_doc
def test_doc_handler_refuses_without_libreoffice(monkeypatch, tmp_path):
    """LibreOffice 不存在时给出清晰的错误提示。"""
    monkeypatch.setattr(
        "docmask.handlers.doc_handler.DocHandler._find_libreoffice_command",
        lambda _self: None,
    )
    monkeypatch.setattr(
        "docmask.handlers.doc_handler.DocHandler._try_pywin32_convert",
        lambda _self, *_args: None,
    )

    handler = DocHandler()
    with pytest.raises(RuntimeError, match="无法转换 .doc 文件"):
        handler._convert_to_docx(str(TEST_DOC), str(tmp_path))


@skip_if_no_libreoffice
@skip_if_no_test_doc
def test_doc_mask_with_duplicate_original_rule_raises_error(tmp_path):
    """重复的原文规则在 Codebook 加载时拒止（审计 P0-12）。"""
    from docmask.core.codebook import CodebookError

    codebook_path = tmp_path / "codebook_dup.txt"
    codebook_path.write_text(
        "假药集团==>MaskA\n假药集团==>MaskB\n", encoding="utf-8"
    )
    with pytest.raises(CodebookError, match="重复"):
        Codebook(str(codebook_path)).load()


# ======================== 残余扫描：脱敏后保存的 DOCX 不含敏感原文 ========================

@skip_if_no_libreoffice
@skip_if_no_test_doc
def test_residual_scan_passes_after_masking(tmp_path):
    """脱敏后残留扫描通过——支持的 OPC 部件中无敏感原文。"""
    masker, _restorer = make_codebook(
        tmp_path,
        "假药集团==>MASKED_GROUP\n"
        "李加薪==>MASKED_NAME\n"
        "土城头==>MASKED_PROJ\n"
        "白蛇河==>MASKED_PROJ2\n",
    )
    handler = DocHandler()

    # 这不应抛出 RuntimeError
    masked_path, count, _coverage = handler.mask(
        str(TEST_DOC), masker, output_path=str(tmp_path / "out.doc"),
    )

    assert count > 0
    assert handler._is_valid_docx(masked_path)
    # 残留扫描通过 — 没有因残留校验崩溃


# ======================== 脱敏词冲突检测 ========================

@skip_if_no_libreoffice
@skip_if_no_test_doc
def test_conflict_precheck_runs_on_doc(tmp_path):
    """DOC 脱敏时执行冲突预检。"""
    masker, _restorer = make_codebook(
        tmp_path,
        "假药集团==>某集团公司\n"  # 脱敏词不与文档正文冲突
    )
    handler = DocHandler()

    masked_path, count, _coverage = handler.mask(
        str(TEST_DOC), masker, output_path=str(tmp_path / "out.doc"),
    )
    assert count > 0

    # 再验证如果脱敏词与文档正文冲突会拒止
    # "投资并购" 出现在 P0 中
    masker2, _ = make_codebook(tmp_path, "假药集团==>投资并购\n")
    from docmask.core.masker import MaskConflictError
    with pytest.raises(MaskConflictError, match="文档中已存在脱敏词"):
        handler.mask(str(TEST_DOC), masker2, output_path=str(tmp_path / "out2.doc"))
