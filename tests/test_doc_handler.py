"""DOC 转换链路的隔离回归测试。"""
from __future__ import annotations

import shutil
import sys
import types
from pathlib import Path
from urllib.parse import unquote, urlparse

from docx import Document

from docmask.core.codebook import Codebook
from docmask.core.masker import Masker
from docmask.core.restorer import Restorer
from docmask.handlers.doc_handler import DocHandler
from tests.word_fixture_factory import create_complex_word_fixture


def _engines(tmp_path: Path):
    codebook_path = tmp_path / "codebook.txt"
    codebook_path.write_text("张三==>匿名客户001\n", encoding="utf-8")
    codebook = Codebook(str(codebook_path))
    codebook.load()
    return Masker(codebook), Restorer(codebook)


def test_convert_returns_converter_output_instead_of_new_empty_file(tmp_path, monkeypatch):
    source_docx = create_complex_word_fixture(tmp_path / "source.docx")
    legacy_doc = tmp_path / "source.doc"
    legacy_doc.write_bytes(b"isolated test placeholder")
    handler = DocHandler()

    def fake_word_convert(_input_path, output_path):
        shutil.copy2(source_docx, output_path)
        return output_path

    monkeypatch.setattr(handler, "_try_pywin32_convert", fake_word_convert)
    monkeypatch.setattr(
        handler,
        "_try_libreoffice_convert",
        lambda *_args: (_ for _ in ()).throw(AssertionError("不应调用 LibreOffice")),
    )

    conversion_dir = tmp_path / "conversion"
    conversion_dir.mkdir()
    converted = Path(handler._convert_to_docx(str(legacy_doc), str(conversion_dir)))

    assert converted == conversion_dir / "source.docx"
    assert converted.read_bytes() == source_docx.read_bytes()
    assert handler._is_valid_docx(converted)


def test_doc_mask_restore_outputs_valid_docx_and_cleans_temp_dir(tmp_path, monkeypatch):
    source_docx = create_complex_word_fixture(tmp_path / "source.docx")
    legacy_doc = tmp_path / "source.doc"
    legacy_doc.write_bytes(b"isolated test placeholder")
    masker, restorer = _engines(tmp_path)
    handler = DocHandler()
    conversion_dirs = []
    conversion_source = {"path": source_docx}

    def fake_convert(_input_path, temp_dir):
        conversion_dirs.append(Path(temp_dir))
        converted = Path(temp_dir) / "source.docx"
        shutil.copy2(conversion_source["path"], converted)
        return str(converted)

    monkeypatch.setattr(handler, "_convert_to_docx", fake_convert)

    masked_path, count, _coverage = handler.mask(
        str(legacy_doc), masker, output_path=str(tmp_path / "masked.doc"),
    )
    conversion_source["path"] = Path(masked_path)
    restored_path, restored_count = handler.restore(
        str(legacy_doc), restorer, output_path=str(tmp_path / "restored.doc"),
    )

    assert Path(masked_path).suffix == ".docx"
    assert Path(restored_path).suffix == ".docx"
    assert handler._is_valid_docx(masked_path)
    assert handler._is_valid_docx(restored_path)
    assert count > 0
    assert restored_count == count
    assert conversion_dirs and all(not path.exists() for path in conversion_dirs)

    masked_document = Document(masked_path)
    assert "匿名客户001" in "\n".join(p.text for p in masked_document.paragraphs)
    restored_document = Document(restored_path)
    assert "张三" in "\n".join(p.text for p in restored_document.paragraphs)


def test_word_com_resources_are_closed_on_conversion_failure(tmp_path, monkeypatch):
    lifecycle = []

    class FakeDocument:
        def SaveAs2(self, *_args, **_kwargs):
            lifecycle.append("save")
            raise RuntimeError("simulated save failure")

        def Close(self, **_kwargs):
            lifecycle.append("close")

    class FakeDocuments:
        def Open(self, *_args, **_kwargs):
            lifecycle.append("open")
            return FakeDocument()

    class FakeWord:
        def __init__(self):
            self.Documents = FakeDocuments()
            self.Visible = True
            self.DisplayAlerts = 1

        def Quit(self):
            lifecycle.append("quit")

    fake_pythoncom = types.ModuleType("pythoncom")
    fake_pythoncom.CoInitialize = lambda: lifecycle.append("coinitialize")
    fake_pythoncom.CoUninitialize = lambda: lifecycle.append("couninitialize")
    fake_client = types.ModuleType("win32com.client")
    fake_client.DispatchEx = lambda _name: FakeWord()
    fake_win32com = types.ModuleType("win32com")
    fake_win32com.client = fake_client

    monkeypatch.setitem(sys.modules, "pythoncom", fake_pythoncom)
    monkeypatch.setitem(sys.modules, "win32com", fake_win32com)
    monkeypatch.setitem(sys.modules, "win32com.client", fake_client)

    input_path = tmp_path / "source.doc"
    input_path.write_bytes(b"isolated test placeholder")
    output_path = tmp_path / "source.docx"

    result = DocHandler()._try_pywin32_convert(str(input_path), output_path)

    assert result is None
    assert lifecycle == [
        "coinitialize", "open", "save", "close", "quit", "couninitialize",
    ]


def test_libreoffice_uses_temporary_profile_and_cleans_it(tmp_path, monkeypatch):
    source_docx = create_complex_word_fixture(tmp_path / "source_fixture.docx")
    legacy_doc = tmp_path / "source.doc"
    legacy_doc.write_bytes(b"isolated test placeholder")
    output_path = tmp_path / "source.docx"
    observed_profile = {"path": None}
    handler = DocHandler()

    monkeypatch.setattr(handler, "_find_libreoffice_command", lambda: "soffice.com")

    def fake_run(args, **_kwargs):
        profile_arg = next(
            arg for arg in args if arg.startswith("-env:UserInstallation=")
        )
        profile_uri = profile_arg.split("=", 1)[1]
        parsed = urlparse(profile_uri)
        profile_path = Path(unquote(parsed.path))
        observed_profile["path"] = profile_path
        assert profile_path.is_dir()
        shutil.copy2(source_docx, output_path)
        return types.SimpleNamespace(returncode=0, stdout="", stderr="")

    monkeypatch.setattr("docmask.handlers.doc_handler.subprocess.run", fake_run)

    converted = handler._try_libreoffice_convert(str(legacy_doc), output_path)

    assert converted == output_path
    assert observed_profile["path"] is not None
    assert not observed_profile["path"].exists()
