"""为 DOC/DOCX 集成测试生成完全隔离的确定性 Word 文稿。"""
from __future__ import annotations

import base64
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.shared import Inches


_ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _append_hyperlink(paragraph, first: str, second: str) -> None:
    relationship_id = paragraph.part.relate_to(
        "https://example.invalid/docmask",
        RT.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    for text, bold in ((first, True), (second, False)):
        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")
        if bold:
            run_properties.append(OxmlElement("w:b"))
        run.append(run_properties)
        text_node = OxmlElement("w:t")
        text_node.text = text
        run.append(text_node)
        hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _note_xml(kind: str, text: str) -> bytes:
    plural = "footnotes" if kind == "footnote" else "endnotes"
    singular = kind
    ref_name = "footnoteRef" if kind == "footnote" else "endnoteRef"
    separator_name = "separator"
    continuation_name = "continuationSeparator"
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:{plural} xmlns:w="{namespace}">
  <w:{singular} w:type="separator" w:id="-1"><w:p><w:r><w:{separator_name}/></w:r></w:p></w:{singular}>
  <w:{singular} w:type="continuationSeparator" w:id="0"><w:p><w:r><w:{continuation_name}/></w:r></w:p></w:{singular}>
  <w:{singular} w:id="1"><w:p><w:r><w:{ref_name}/></w:r><w:r><w:t>{text}</w:t></w:r></w:p></w:{singular}>
</w:{plural}>'''.encode("utf-8")


def _add_note_part(document: Document, kind: str, text: str) -> None:
    if kind == "footnote":
        partname = PackURI("/word/footnotes.xml")
        content_type = CT.WML_FOOTNOTES
        relationship_type = RT.FOOTNOTES
        reference_name = "w:footnoteReference"
    else:
        partname = PackURI("/word/endnotes.xml")
        content_type = CT.WML_ENDNOTES
        relationship_type = RT.ENDNOTES
        reference_name = "w:endnoteReference"

    note_part = Part(
        partname,
        content_type,
        _note_xml(kind, text),
        document.part.package,
    )
    document.part.relate_to(note_part, relationship_type)

    marker_run = document.add_paragraph(f"{kind} marker: ").add_run()
    reference = OxmlElement(reference_name)
    reference.set(qn("w:id"), "1")
    marker_run._r.append(reference)


def _append_field_with_sensitive_display(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction.text = " MERGEFIELD CustomerName "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "张三"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, display, end):
        run._r.append(element)


def create_complex_word_fixture(path: str | Path) -> Path:
    """创建包含跨 Run、超链接、脚注、尾注和非文本节点的 DOCX。"""
    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.core_properties.title = "DocMask 隔离集成测试文稿"
    document.core_properties.author = "测试生成器"
    document.add_heading("DocMask Word 集成测试", level=1)

    paragraph = document.add_paragraph("普通正文：")
    first = paragraph.add_run("张")
    first.bold = True
    second = paragraph.add_run("三")
    second.italic = True
    controls = paragraph.add_run()
    controls.add_tab()
    controls.add_text("Tab 后文字保持")
    controls.add_break()
    controls.add_text("换行后文字保持")

    link_paragraph = document.add_paragraph("超链接显示文本：")
    _append_hyperlink(link_paragraph, "张", "三")

    field_paragraph = document.add_paragraph("域显示文本：")
    _append_field_with_sensitive_display(field_paragraph)

    drawing_paragraph = document.add_paragraph("绘图节点：")
    drawing_run = drawing_paragraph.add_run()
    drawing_run.add_picture(BytesIO(_ONE_PIXEL_PNG), width=Inches(0.2))
    drawing_paragraph.add_run(" 张三")

    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "张三"

    header = document.sections[0].header.paragraphs[0]
    header.text = "页眉 张三"

    _add_note_part(document, "footnote", "脚注中的张三")
    _add_note_part(document, "endnote", "尾注中的张三")

    document.save(output_path)
    return output_path
