"""集成测试：完整流程"""
import os
import tempfile
import shutil
import pytest
from docx import Document
from docmask.core.codebook import Codebook
from docmask.core.masker import Masker
from docmask.core.restorer import Restorer
from docmask.handlers.txt_handler import TxtHandler
from docmask.handlers.docx_handler import DocxHandler

DATA_DIR = os.path.join(os.path.dirname(__file__), "test_data")


@pytest.fixture
def codebook():
    cb = Codebook(os.path.join(DATA_DIR, "sample_codebook.txt"))
    cb.load()
    return cb


class TestIntegration:
    """集成测试"""

    def test_txt_full_flow(self, codebook):
        """TXT: 原文→脱敏→恢复→对比"""
        masker = Masker(codebook)
        restorer = Restorer(codebook)
        handler = TxtHandler()

        original = "张三和李明在某某科技有限公司工作"
        with tempfile.NamedTemporaryFile(
            suffix=".txt", mode="w", delete=False, encoding="utf-8"
        ) as f:
            f.write(original)
            tmp_in = f.name

        try:
            # 脱敏
            masked_path, count, _ = handler.mask(tmp_in, masker)
            # 恢复
            restored_path, _ = handler.restore(masked_path, restorer)
            # 对比
            with open(restored_path, "r", encoding="utf-8") as f:
                restored = f.read()
            assert restored == original
        finally:
            for p in [tmp_in, masked_path, restored_path]:
                if os.path.exists(p):
                    os.remove(p)

    def test_docx_full_flow(self, codebook):
        """DOCX: 原文→脱敏→恢复→对比"""
        masker = Masker(codebook)
        restorer = Restorer(codebook)
        handler = DocxHandler()

        doc = Document()
        doc.core_properties.author = "张三"
        doc.add_paragraph("张三在某某科技有限公司工作，同事李明")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_in = f.name

        try:
            # 脱敏
            masked_path, count, _ = handler.mask(tmp_in, masker)
            assert count > 0
            # 恢复
            restored_path, _ = handler.restore(masked_path, restorer)
            result_doc = Document(restored_path)
            texts = [p.text for p in result_doc.paragraphs]
            combined = "".join(texts)
            assert "张三" in combined
            assert "某某科技有限公司" in combined
            assert "李明" in combined
        finally:
            for p in [tmp_in, masked_path, restored_path]:
                if os.path.exists(p):
                    os.remove(p)

    def test_codebook_validation(self):
        """密码本校验：错误密码本"""
        cb = Codebook(os.path.join(DATA_DIR, "codebook_conflict.txt"))
        cb.load()
        messages = cb.validate()
        errors = [m for m in messages if m.startswith("ERROR")]
        assert len(errors) > 0

    def test_batch_txt(self, codebook):
        """批量处理：目录下多个 TXT"""
        masker = Masker(codebook)
        handler = TxtHandler()

        tmp_dir = tempfile.mkdtemp()
        try:
            # 创建两个文件
            for i, name in enumerate(["张三", "李明"]):
                with open(os.path.join(tmp_dir, f"test_{i}.txt"), "w", encoding="utf-8") as f:
                    f.write(name)

            # 批量处理
            files = [os.path.join(tmp_dir, f) for f in os.listdir(tmp_dir)]
            for fp in files:
                if fp.endswith(".txt"):
                    handler.mask(fp, masker)

            # 验证
            for f in os.listdir(tmp_dir):
                if f.endswith("_desensitized.txt"):
                    with open(os.path.join(tmp_dir, f), "r", encoding="utf-8") as fh:
                        content = fh.read()
                    assert "张三" not in content
                    assert "李明" not in content
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)