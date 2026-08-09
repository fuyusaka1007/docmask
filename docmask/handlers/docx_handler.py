"""DOCX 文件读写与脱敏/恢复。

仅对能力矩阵中声明支持的 OPC 部件做强保证；发现其他部件含疑似敏感文本时
返回警告，避免用“覆盖全元素”掩盖 OOXML 的能力边界。
"""
import logging
import re
import zipfile
from difflib import SequenceMatcher
from pathlib import Path
from typing import Optional

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree
from docx.oxml.ns import qn

from docmask.core.masker import Masker, MaskConflictError
from docmask.core.restorer import Restorer
from docmask.utils.file_utils import generate_output_path, staged_output_path
from docmask.config import DESENSITIZED_SUFFIX, RESTORED_SUFFIX
from docmask.handlers.base import (
    CancelToken,
    ProgressCallback,
    check_cancel,
    report_progress,
)

logger = logging.getLogger(__name__)

# XML namespaces
NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
    "v": "urn:schemas-microsoft-com:vml",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}

# ===== 扩展隐私表面（统一选择器：采集/脱敏/恢复/校验共用） =====
# 支持部件内 w:t/w:delText 之外的可安全替换文本节点（element localname 集合）。
# w:instrText —— 域代码指令文本（如 MERGEFIELD CustomerName）。
_EXTENDED_TEXT_LOCALNAMES = {"instrText"}

# 支持部件内属性级文本：element localname -> [(attribute localname, namespace)]
#   namespace 为 "w" 表示属性在 w: 命名空间下；None 表示无命名空间属性。
#   w:sdtPr/w:tag/@w:val          —— 内容控件标签
#   w:bookmarkStart/@w:name       —— 书签名称（有 XML Name 命名约束）
#   wp:docPr/@descr / @title      —— 绘图/图片替代文本
_EXTENDED_ATTRIBUTE_SURFACE: dict[str, list[tuple[str, Optional[str]]]] = {
    "tag": [("val", "w")],
    "bookmarkStart": [("name", "w")],
    "docPr": [("descr", None), ("title", None)],
}

# 书签名称必须符合 XML Name 规范：字母/下划线开头，仅含字母数字下划线连字符句点。
# 脱敏词不符合此规范时 fail closed（阻止输出），避免静默保留敏感内容。
_BOOKMARK_NAME_RE = re.compile(r"^[A-Za-z_][A-Za-z0-9_.\-]*$")


class DocxHandler:
    """DOCX 文件读写与脱敏/恢复，保留格式"""

    OPC_CAPABILITY_MATRIX = {
        "word/document.xml": "supported",
        "word/header*.xml": "supported",
        "word/footer*.xml": "supported",
        "word/footnotes.xml": "supported",
        "word/endnotes.xml": "supported",
        "word/comments*.xml": "supported",
        "docProps/core.xml": "supported",
        "docProps/app.xml": "supported-limited",
        "docProps/custom.xml": "supported-limited",
        "word/charts/*": "warn-only",
        "word/diagrams/*": "warn-only",
        "word/embeddings/*": "warn-only",
        "external hyperlink targets": "blocked-by-default",
    }
    # 扩展隐私表面：支持部件内 w:t/w:delText 之外的文本节点与属性
    EXTENDED_PRIVACY_SURFACE = {
        "text_nodes": sorted(_EXTENDED_TEXT_LOCALNAMES),
        "attributes": {
            local: [(a, ns) for a, ns in attrs]
            for local, attrs in _EXTENDED_ATTRIBUTE_SURFACE.items()
        },
    }
    CORE_TEXT_FIELDS = {
        "creator", "lastModifiedBy", "title", "subject", "category",
        "description", "contentStatus", "identifier", "keywords",
        "language", "version",
    }

    def __init__(self):
        self._total_count = 0
        self.last_warnings: list[str] = []
        # 外部关系目标安全策略：True（默认，安全模式）含敏感原文时阻止输出；
        # False（兼容模式）仅告警。
        self.strict_external_targets = True

    # ======================== 安全 XML 文本写回 ========================

    @staticmethod
    def _iter_direct_text_groups(parent_element):
        """按文档顺序返回连续的 w:t 组，遇到非文本节点即切断。

        只遍历 parent 的直接 w:r 子元素。这样普通段落不会误处理
        w:hyperlink 内的 Run；Tab、换行、域、绘图等节点也不会被跨越。
        """
        transparent_containers = {
            qn("w:ins"), qn("w:del"), qn("w:smartTag"),
            qn("w:customXml"), qn("w:sdt"), qn("w:sdtContent"),
        }

        def walk(container):
            group = []
            for child in container:
                if child.tag == qn("w:r"):
                    for run_child in child:
                        if run_child.tag == qn("w:rPr"):
                            continue
                        if run_child.tag in (qn("w:t"), qn("w:delText")):
                            group.append(run_child)
                        elif group:
                            yield group
                            group = []
                    continue

                if group:
                    yield group
                    group = []
                if child.tag in transparent_containers:
                    yield from walk(child)
            if group:
                yield group

        yield from walk(parent_element)

    @staticmethod
    def _set_text_node(node, value: str) -> None:
        """设置 w:t/w:delText 内容，并正确维护 xml:space。"""
        node.text = value
        xml_space = "{http://www.w3.org/XML/1998/namespace}space"
        if value[:1].isspace() or value[-1:].isspace():
            node.set(xml_space, "preserve")
        else:
            node.attrib.pop(xml_space, None)

    def _write_back_text_nodes(self, text_nodes, replaced_text: str) -> None:
        """只更新 w:t 节点，保留 Run 中所有非文本 XML 子节点。"""
        original_texts = [node.text or "" for node in text_nodes]
        original_text = "".join(original_texts)
        if original_text == replaced_text:
            return

        offsets = []
        cursor = 0
        for value in original_texts:
            offsets.append((cursor, cursor + len(value)))
            cursor += len(value)

        buffers = ["" for _ in text_nodes]

        def owner_index(position: int) -> int:
            if not text_nodes:
                return 0
            for index, (start, end) in enumerate(offsets):
                if start <= position < end:
                    return index
                if position == start and start == end:
                    return index
            if position >= len(original_text):
                return len(text_nodes) - 1
            for index, (start, _end) in enumerate(offsets):
                if start >= position:
                    return index
            return len(text_nodes) - 1

        matcher = SequenceMatcher(None, original_text, replaced_text, autojunk=False)
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == "equal":
                for index, (start, end) in enumerate(offsets):
                    overlap_start = max(start, i1)
                    overlap_end = min(end, i2)
                    if overlap_start < overlap_end:
                        buffers[index] += original_text[overlap_start:overlap_end]
            elif tag in ("replace", "insert"):
                buffers[owner_index(i1)] += replaced_text[j1:j2]
            # delete 不写入任何内容

        for node, value in zip(text_nodes, buffers):
            self._set_text_node(node, value)

    def _mask_direct_runs(self, parent_element, masker: Masker) -> tuple[int, dict]:
        total = 0
        hits: dict[str, int] = {}
        for text_nodes in self._iter_direct_text_groups(parent_element):
            original = "".join(node.text or "" for node in text_nodes)
            if not original:
                continue
            masked, count, chunk_hits = masker.mask_text(original)
            if count:
                self._write_back_text_nodes(text_nodes, masked)
                total += count
                for key, value in chunk_hits.items():
                    hits[key] = hits.get(key, 0) + value
        return total, hits

    def _restore_direct_runs(self, parent_element, restorer: Restorer) -> int:
        total = 0
        for text_nodes in self._iter_direct_text_groups(parent_element):
            original = "".join(node.text or "" for node in text_nodes)
            if not original:
                continue
            restored, count = restorer.restore_text(original)
            if count:
                self._write_back_text_nodes(text_nodes, restored)
                total += count
        return total

    @staticmethod
    def _header_footer_roots(doc: Document):
        """返回包中已存在的页眉页脚 XML 根，不创建新的空 part。"""
        seen_parts = set()
        for relationship in doc.part.rels.values():
            if relationship.reltype not in (RT.HEADER, RT.FOOTER):
                continue
            story_part = relationship.target_part
            part_key = str(story_part.partname)
            if part_key in seen_parts:
                continue
            seen_parts.add(part_key)
            yield story_part.element

    @staticmethod
    def _note_parts(doc: Document):
        """按 relationship type 查找脚注和尾注 part。"""
        supported = {RT.FOOTNOTES: "脚注", RT.ENDNOTES: "尾注"}
        for relationship in doc.part.rels.values():
            label = supported.get(relationship.reltype)
            if label is not None:
                yield label, relationship.target_part

    # ======================== 脱敏入口 ========================

    def mask(
        self,
        input_path: str,
        masker: Masker,
        output_path: Optional[str] = None,
        progress_callback: Optional[ProgressCallback] = None,
        cancel_token: Optional[CancelToken] = None,
    ) -> tuple[str, int, dict]:
        """
        脱敏 docx 文件（正文、表格、页眉页脚、文本框、脚注尾注、超链接、元数据）
        返回 (输出路径, 替换次数, 覆盖率统计)

        进度步骤共 9 步：冲突预检 + 支持的 OPC 文本部件
        """
        TOTAL_STEPS = 9
        self._total_count = 0
        self.last_warnings = []
        accumulated_hits: dict[str, int] = {}
        doc = Document(input_path)

        # 步骤 1: 文档级脱敏词冲突预检
        report_progress(progress_callback, 0, TOTAL_STEPS, "正在执行冲突预检...")
        check_cancel(cancel_token)
        all_text = self._collect_all_text(doc)
        conflicts = masker.precheck_conflict(all_text)
        if conflicts:
            raise MaskConflictError(MaskConflictError.format(conflicts))
        # 脱敏前快照：收集所有部件的原始文本，用于保存后区分"未被替换的原文"
        # 和"脱敏词中的字符"，避免脱敏词自污染导致的误报。
        original_texts = self._snapshot_all_texts(doc)
        report_progress(progress_callback, 1, TOTAL_STEPS, "冲突预检完成")

        count = 0

        def _accumulate(added_count: int, hits: dict[str, int]):
            nonlocal count
            count += added_count
            for k, v in hits.items():
                accumulated_hits[k] = accumulated_hits.get(k, 0) + v

        # 步骤 2-8: 逐元素脱敏
        _accumulate(*self._mask_paragraphs(doc.paragraphs, masker))
        _accumulate(*self._mask_sdt_blocks(doc, masker))
        report_progress(progress_callback, 2, TOTAL_STEPS, "正文段落脱敏完成")
        check_cancel(cancel_token)

        _accumulate(*self._mask_tables(doc, masker))
        report_progress(progress_callback, 3, TOTAL_STEPS, "表格脱敏完成")
        check_cancel(cancel_token)

        _accumulate(*self._mask_headers_footers(doc, masker))
        report_progress(progress_callback, 4, TOTAL_STEPS, "页眉页脚脱敏完成")
        check_cancel(cancel_token)

        _accumulate(*self._mask_textboxes(doc, masker))
        report_progress(progress_callback, 5, TOTAL_STEPS, "文本框脱敏完成")
        check_cancel(cancel_token)

        _accumulate(*self._mask_footnotes(doc, masker))
        report_progress(progress_callback, 6, TOTAL_STEPS, "脚注尾注脱敏完成")
        check_cancel(cancel_token)

        _accumulate(*self._mask_hyperlinks(doc, masker))
        report_progress(progress_callback, 7, TOTAL_STEPS, "超链接脱敏完成")
        check_cancel(cancel_token)

        _accumulate(*self._mask_metadata(doc, masker))
        report_progress(progress_callback, 8, TOTAL_STEPS, "元数据脱敏完成")

        _accumulate(*self._mask_auxiliary_parts(doc, masker))
        _accumulate(*self._mask_extended_surface(doc, masker))
        self.last_warnings.extend(
            self._scan_unsupported_parts(doc, masker, original_texts)
        )
        report_progress(progress_callback, 9, TOTAL_STEPS, "OPC 扩展部件检查完成")

        self._total_count = count

        if output_path is None:
            output_path = generate_output_path(input_path, suffix=DESENSITIZED_SUFFIX)

        with staged_output_path(output_path) as temp_path:
            doc.save(temp_path)
            self._assert_no_supported_residuals(temp_path, masker, original_texts)
        logger.info(f"脱敏完成: {input_path} -> {output_path}, 替换 {count} 处")
        return output_path, count, dict(accumulated_hits)

    # ======================== 恢复入口 ========================

    def restore(
        self,
        input_path: str,
        restorer: Restorer,
        output_path: Optional[str] = None,
        progress_callback: Optional[ProgressCallback] = None,
        cancel_token: Optional[CancelToken] = None,
    ) -> tuple[str, int]:
        """
        恢复 docx 文件（覆盖所有与脱敏相同的元素）
        返回 (输出路径, 替换次数)

        进度步骤共 8 步：支持的 OPC 文本部件
        """
        TOTAL_STEPS = 8
        self._total_count = 0
        self.last_warnings = []
        doc = Document(input_path)

        count = 0

        count += self._restore_paragraphs(doc.paragraphs, restorer)
        count += self._restore_sdt_blocks(doc, restorer)
        report_progress(progress_callback, 1, TOTAL_STEPS, "正文段落恢复完成")
        check_cancel(cancel_token)

        count += self._restore_tables(doc, restorer)
        report_progress(progress_callback, 2, TOTAL_STEPS, "表格恢复完成")
        check_cancel(cancel_token)

        count += self._restore_headers_footers(doc, restorer)
        report_progress(progress_callback, 3, TOTAL_STEPS, "页眉页脚恢复完成")
        check_cancel(cancel_token)

        count += self._restore_textboxes(doc, restorer)
        report_progress(progress_callback, 4, TOTAL_STEPS, "文本框恢复完成")
        check_cancel(cancel_token)

        count += self._restore_footnotes(doc, restorer)
        report_progress(progress_callback, 5, TOTAL_STEPS, "脚注尾注恢复完成")
        check_cancel(cancel_token)

        count += self._restore_hyperlinks(doc, restorer)
        report_progress(progress_callback, 6, TOTAL_STEPS, "超链接恢复完成")
        check_cancel(cancel_token)

        count += self._restore_metadata(doc, restorer)
        report_progress(progress_callback, 7, TOTAL_STEPS, "元数据恢复完成")

        count += self._restore_auxiliary_parts(doc, restorer)
        count += self._restore_extended_surface(doc, restorer)
        report_progress(progress_callback, 8, TOTAL_STEPS, "OPC 扩展部件恢复完成")

        self._total_count = count

        if output_path is None:
            output_path = generate_output_path(input_path, suffix=RESTORED_SUFFIX)

        with staged_output_path(output_path) as temp_path:
            doc.save(temp_path)
            # A-17: 恢复后残留校验，与脱敏阶段对称
            self._assert_no_masked_word_residuals(temp_path, restorer)
        logger.info(f"恢复完成: {input_path} -> {output_path}, 替换 {count} 处")
        return output_path, count

    # ======================== 文本收集（冲突预检用） ========================

    @classmethod
    def _collect_grouped_texts(cls, root) -> list[str]:
        """A-05: 使用与 _mask_direct_runs 相同的文本分组逻辑收集文本。

        确保冲突预检看到的文本与实际脱敏处理的文本坐标一致，
        避免跨 Run 拆分的脱敏词（如 MAS + KED）在预检中漏检。
        """
        texts: list[str] = []
        for para in root.findall(".//w:p", NSMAP):
            for text_nodes in cls._iter_direct_text_groups(para):
                text = "".join(node.text or "" for node in text_nodes)
                if text:
                    texts.append(text)
            for hyperlink in para.findall("w:hyperlink", NSMAP):
                for text_nodes in cls._iter_direct_text_groups(hyperlink):
                    text = "".join(node.text or "" for node in text_nodes)
                    if text:
                        texts.append(text)
        return texts

    def _collect_all_text(self, doc: Document) -> str:
        """收集文档中所有文本内容，用于冲突预检

        A-05: 使用与脱敏阶段相同的文本分组逻辑，确保跨 Run 拆分的
        脱敏词在预检中也能被检测到。
        """
        parts: list[str] = []

        # 正文（含表格、文本框、SDT 等 body 内所有段落）
        parts.extend(self._collect_grouped_texts(doc.element.body))

        # 页眉页脚（含其中的表格、文本框与超链接）
        for root in self._header_footer_roots(doc):
            parts.extend(self._collect_grouped_texts(root))

        # 脚注和尾注
        for _label, note_part in self._note_parts(doc):
            try:
                root = etree.fromstring(note_part.blob)
            except (etree.XMLSyntaxError, ValueError):
                continue
            parts.extend(self._collect_grouped_texts(root))

        # 批注、扩展属性和自定义属性
        for part in doc.part.package.parts:
            partname = str(part.partname).lstrip("/")
            is_comments = partname.startswith("word/comments") and partname.endswith(".xml")
            if is_comments or partname in {"docProps/app.xml", "docProps/custom.xml"}:
                try:
                    root = etree.fromstring(part.blob)
                    if is_comments:
                        parts.extend(self._collect_grouped_texts(root))
                        for comment in root.findall(".//w:comment", NSMAP):
                            for attribute in (qn("w:author"), qn("w:initials")):
                                value = comment.get(attribute)
                                if value:
                                    parts.append(value)
                    parts.extend(
                        node.text for node in self._auxiliary_part_nodes(root, partname)
                        if node.text
                    )
                except (etree.XMLSyntaxError, ValueError):
                    continue

        # 扩展隐私表面（域代码指令、SDT 标签、书签名、图片替代文本）——用于冲突预检
        parts.extend(self._iter_extended_text_values(doc.element))
        parts.extend(self._iter_extended_attr_values(doc.element))
        for root in self._header_footer_roots(doc):
            parts.extend(self._iter_extended_text_values(root))
            parts.extend(self._iter_extended_attr_values(root))
        for _label, note_part in self._note_parts(doc):
            try:
                root = etree.fromstring(note_part.blob)
            except (etree.XMLSyntaxError, ValueError):
                continue
            parts.extend(self._iter_extended_text_values(root))
            parts.extend(self._iter_extended_attr_values(root))
        for part in doc.part.package.parts:
            partname = str(part.partname).lstrip("/")
            if not (partname.startswith("word/comments") and partname.endswith(".xml")):
                continue
            try:
                root = etree.fromstring(part.blob)
            except (etree.XMLSyntaxError, ValueError):
                continue
            parts.extend(self._iter_extended_text_values(root))
            parts.extend(self._iter_extended_attr_values(root))

        # 元数据
        props = doc.core_properties
        for name in (
            "author", "last_modified_by", "title", "subject", "category",
            "comments", "content_status", "identifier", "keywords",
            "language", "version",
        ):
            value = getattr(props, name, None)
            if value:
                parts.append(value)

        return "\n".join(parts)

    # ======================== 段落处理（脱敏） ========================

    def _mask_paragraphs(self, paragraphs, masker: Masker) -> tuple[int, dict]:
        """脱敏段落列表（正文/表格/页眉页脚共用），返回 (替换次数, 命中统计)"""
        total = 0
        hits: dict[str, int] = {}
        for para in paragraphs:
            count, chunk_hits = self._mask_direct_runs(para._p, masker)
            total += count
            for key, value in chunk_hits.items():
                hits[key] = hits.get(key, 0) + value

        return total, hits

    def _write_back_to_runs(self, runs, original_texts: list[str], replaced_text: str) -> None:
        """兼容旧调用：只修改 Run 内的 w:t，不重建 Run。"""
        text_nodes = []
        for run in runs:
            text_nodes.extend(run._r.findall("w:t", NSMAP))
        if text_nodes:
            self._write_back_text_nodes(text_nodes, replaced_text)

    # ======================== 表格处理 ========================

    def _mask_tables(self, doc: Document, masker: Masker) -> tuple[int, dict]:
        """脱敏表格中的文本"""
        total = 0
        hits: dict[str, int] = {}
        def walk(table):
            nonlocal total
            for row in table.rows:
                for cell in row.cells:
                    c, h = self._mask_paragraphs(cell.paragraphs, masker)
                    total += c
                    for k, v in h.items():
                        hits[k] = hits.get(k, 0) + v
                    for nested in cell.tables:
                        walk(nested)
                    # 补齐 cell 内的 sdt（cell.paragraphs/cell.tables 不含 sdt 内内容）
                    for sdt_elem in cell._tc.findall(qn("w:sdt")):
                        c, h = self._mask_sdt_content(sdt_elem, masker)
                        total += c
                        for k, v in h.items():
                            hits[k] = hits.get(k, 0) + v

        for table in doc.tables:
            walk(table)
        return total, hits

    # ======================== SDT 内容控件处理 ========================

    def _mask_sdt_blocks(self, doc: Document, masker: Masker) -> tuple[int, dict]:
        """脱敏 body 中 w:sdt 容器内的段落和表格。

        doc.paragraphs 和 doc.tables 只返回 body 的直接子段落/表格，
        不包含被 w:sdt/w:sdtContent 包裹的内容。此方法补齐该盲区，
        递归处理 sdt 内的段落、表格和嵌套 sdt。

        不会与 _mask_textboxes/_mask_hyperlinks 重叠：_mask_direct_runs
        只处理段落的直接 Run，不递归进 txbx/hyperlink。
        """
        total = 0
        hits: dict[str, int] = {}
        body = doc.element.body
        for sdt_elem in body.findall(qn("w:sdt")):
            c, h = self._mask_sdt_content(sdt_elem, masker)
            total += c
            for k, v in h.items():
                hits[k] = hits.get(k, 0) + v
        return total, hits

    def _mask_sdt_content(self, sdt_elem, masker: Masker) -> tuple[int, dict]:
        """递归脱敏 w:sdt 内的段落、表格和嵌套 sdt。"""
        total = 0
        hits: dict[str, int] = {}
        content = sdt_elem.find(qn("w:sdtContent"))
        if content is None:
            return total, hits
        for child in content:
            if child.tag == qn("w:p"):
                c, h = self._mask_direct_runs(child, masker)
            elif child.tag == qn("w:tbl"):
                c, h = self._mask_table_xml(child, masker)
            elif child.tag == qn("w:sdt"):
                c, h = self._mask_sdt_content(child, masker)
            else:
                continue
            total += c
            for k, v in h.items():
                hits[k] = hits.get(k, 0) + v
        return total, hits

    def _mask_table_xml(self, table_elem, masker: Masker) -> tuple[int, dict]:
        """脱敏 XML 表格元素中的文本（递归嵌套表格和 sdt）。

        与 _mask_tables 不同，此方法直接操作 XML 元素，
        用于 sdt 内的表格（doc.tables 不返回 sdt 内的表格）。
        """
        total = 0
        hits: dict[str, int] = {}
        for row in table_elem.findall(qn("w:tr")):
            for cell in row.findall(qn("w:tc")):
                for para_elem in cell.findall(qn("w:p")):
                    c, h = self._mask_direct_runs(para_elem, masker)
                    total += c
                    for k, v in h.items():
                        hits[k] = hits.get(k, 0) + v
                for nested_table in cell.findall(qn("w:tbl")):
                    c, h = self._mask_table_xml(nested_table, masker)
                    total += c
                    for k, v in h.items():
                        hits[k] = hits.get(k, 0) + v
                for nested_sdt in cell.findall(qn("w:sdt")):
                    c, h = self._mask_sdt_content(nested_sdt, masker)
                    total += c
                    for k, v in h.items():
                        hits[k] = hits.get(k, 0) + v
        return total, hits

    # ======================== 页眉页脚处理 ========================

    def _mask_headers_footers(self, doc: Document, masker: Masker) -> tuple[int, dict]:
        """脱敏页眉页脚中的文本"""
        total = 0
        hits: dict[str, int] = {}
        for root in self._header_footer_roots(doc):
            count, chunk_hits = self._mask_xml_paragraphs(root, masker)
            total += count
            for key, value in chunk_hits.items():
                hits[key] = hits.get(key, 0) + value
        return total, hits

    # ======================== 文本框处理 ========================

    def _mask_textboxes(self, doc: Document, masker: Masker) -> tuple[int, dict]:
        """脱敏文本框中的文本（操作底层XML）"""
        total = 0
        hits: dict[str, int] = {}
        body = doc.element.body

        # 查找所有 w:txbxContent 元素（文本框内容）
        txbx_elements = body.findall(".//w:txbxContent", NSMAP)
        for txbx in txbx_elements:
            c, h = self._mask_xml_paragraphs(txbx, masker)
            total += c
            for k, v in h.items():
                hits[k] = hits.get(k, 0) + v
        return total, hits

    def _mask_xml_paragraphs(
        self,
        parent_element,
        masker: Masker,
        include_hyperlinks: bool = False,
    ) -> tuple[int, dict]:
        """对 XML 元素中的段落进行脱敏（用于文本框和脚注）"""
        total = 0
        hits: dict[str, int] = {}
        para_elements = parent_element.findall(".//w:p", NSMAP)
        for para_elem in para_elements:
            count, chunk_hits = self._mask_direct_runs(para_elem, masker)
            total += count
            for key, value in chunk_hits.items():
                hits[key] = hits.get(key, 0) + value
            if include_hyperlinks:
                for hyperlink in para_elem.findall("w:hyperlink", NSMAP):
                    count, chunk_hits = self._mask_direct_runs(hyperlink, masker)
                    total += count
                    for key, value in chunk_hits.items():
                        hits[key] = hits.get(key, 0) + value

        return total, hits

    def _write_back_to_xml_runs(self, run_elements, original_texts: list[str], replaced_text: str) -> None:
        """兼容旧调用：安全写回 XML Run 中的全部 w:t。"""
        text_nodes = []
        for run_element in run_elements:
            text_nodes.extend(run_element.findall("w:t", NSMAP))
        if text_nodes:
            self._write_back_text_nodes(text_nodes, replaced_text)

    # ======================== 脚注/尾注处理 ========================

    def _mask_footnotes(self, doc: Document, masker: Masker) -> tuple[int, dict]:
        """脱敏脚注/尾注中的文本"""
        total = 0
        hits: dict[str, int] = {}
        try:
            for _label, note_part in self._note_parts(doc):
                root = etree.fromstring(note_part.blob)
                c, h = self._mask_xml_paragraphs(
                    root, masker, include_hyperlinks=True,
                )
                total += c
                for k, v in h.items():
                    hits[k] = hits.get(k, 0) + v
                note_part._blob = etree.tostring(
                    root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )
        except Exception as e:
            raise RuntimeError(f"脚注/尾注处理失败，任务已停止：{e}") from e

        return total, hits

    # ======================== 超链接处理 ========================

    def _mask_hyperlinks(self, doc: Document, masker: Masker) -> tuple[int, dict]:
        """脱敏超链接显示文本"""
        total = 0
        hits: dict[str, int] = {}
        roots = [doc.element.body, *self._header_footer_roots(doc)]
        for root in roots:
            for hyperlink in root.findall(".//w:hyperlink", NSMAP):
                count, chunk_hits = self._mask_direct_runs(hyperlink, masker)
                total += count
                for key, value in chunk_hits.items():
                    hits[key] = hits.get(key, 0) + value
        return total, hits

    # ======================== 元数据处理 ========================

    def _mask_metadata(self, doc: Document, masker: Masker) -> tuple[int, dict]:
        """脱敏文档元数据（作者、标题、公司等）"""
        total = 0
        hits: dict[str, int] = {}
        props = doc.core_properties

        field_names = (
            "author", "last_modified_by", "title", "subject", "category",
            "comments", "content_status", "identifier", "keywords",
            "language", "version",
        )
        fields = {name: getattr(props, name, None) for name in field_names}

        for field_name, value in fields.items():
            if value:
                masked_text, count, chunk_hits = masker.mask_text(value)
                if count > 0:
                    total += count
                    for k, v in chunk_hits.items():
                        hits[k] = hits.get(k, 0) + v
                    setattr(props, field_name, masked_text)

        return total, hits

    @staticmethod
    def _auxiliary_part_nodes(root, partname: str):
        """返回扩展/自定义属性中允许改写的字符串节点。"""
        if partname == "docProps/app.xml":
            allowed = {"Company", "Manager", "Template"}
            return [node for node in root.iter() if etree.QName(node).localname in allowed]
        if partname == "docProps/custom.xml":
            allowed = {"lpwstr", "lpstr", "bstr"}
            return [node for node in root.iter() if etree.QName(node).localname in allowed]
        return []

    @staticmethod
    def _replace_part_xml(part, root) -> None:
        """兼容 XmlPart 与普通 Part 的 XML 写回。"""
        if hasattr(part, "_element"):
            part._element = root
        else:
            part._blob = etree.tostring(
                root, xml_declaration=True, encoding="UTF-8", standalone=True
            )

    def _mask_auxiliary_parts(self, doc: Document, masker: Masker) -> tuple[int, dict]:
        """处理批注、扩展属性和自定义属性 OPC part。"""
        total = 0
        hits: dict[str, int] = {}
        for part in doc.part.package.parts:
            partname = str(part.partname).lstrip("/")
            is_comments = partname.startswith("word/comments") and partname.endswith(".xml")
            if not is_comments and partname not in {"docProps/app.xml", "docProps/custom.xml"}:
                continue
            try:
                root = etree.fromstring(part.blob)
                if is_comments:
                    count, chunk_hits = self._mask_xml_paragraphs(
                        root, masker, include_hyperlinks=True
                    )
                    for comment in root.findall(".//w:comment", NSMAP):
                        for attribute in (qn("w:author"), qn("w:initials")):
                            value = comment.get(attribute)
                            if not value:
                                continue
                            masked, added, node_hits = masker.mask_text(value)
                            if added:
                                comment.set(attribute, masked)
                                count += added
                                for key, amount in node_hits.items():
                                    chunk_hits[key] = chunk_hits.get(key, 0) + amount
                else:
                    count = 0
                    chunk_hits = {}
                    for node in self._auxiliary_part_nodes(root, partname):
                        if not node.text:
                            continue
                        masked, added, node_hits = masker.mask_text(node.text)
                        if added:
                            node.text = masked
                            count += added
                            for key, value in node_hits.items():
                                chunk_hits[key] = chunk_hits.get(key, 0) + value
                if count:
                    self._replace_part_xml(part, root)
                    total += count
                    for key, value in chunk_hits.items():
                        hits[key] = hits.get(key, 0) + value
            except (etree.XMLSyntaxError, ValueError) as exc:
                raise RuntimeError(f"OPC 部件处理失败：{partname}") from exc
        return total, hits

    @staticmethod
    def _is_supported_part(partname: str) -> bool:
        return (
            partname == "word/document.xml"
            or partname.startswith("word/header")
            or partname.startswith("word/footer")
            or partname in {"word/footnotes.xml", "word/endnotes.xml"}
            or partname.startswith("word/comments")
            or partname in {"docProps/core.xml", "docProps/app.xml", "docProps/custom.xml"}
        )

    # ===== 扩展隐私表面：统一选择器与辅助方法 =====

    @staticmethod
    def _get_extended_attr(element, attr_local: str, attr_ns: Optional[str]) -> Optional[str]:
        """读取扩展隐私表面属性值。attr_ns 为 "w" 时用 w: 命名空间，None 为无命名空间。"""
        if attr_ns == "w":
            return element.get(qn(f"w:{attr_local}"))
        return element.get(attr_local)

    @staticmethod
    def _set_extended_attr(element, attr_local: str, attr_ns: Optional[str], value: str) -> None:
        """写入扩展隐私表面属性值。"""
        if attr_ns == "w":
            element.set(qn(f"w:{attr_local}"), value)
        else:
            element.set(attr_local, value)

    @staticmethod
    def _iter_extended_text_values(root) -> list[str]:
        """返回 root 中所有扩展文本节点（w:instrText 等）的非空文本。"""
        values = []
        for node in root.iter():
            if etree.QName(node).localname in _EXTENDED_TEXT_LOCALNAMES and node.text:
                values.append(node.text)
        return values

    @classmethod
    def _iter_extended_attr_values(cls, root) -> list[str]:
        """返回 root 中所有扩展属性节点（SDT tag/bookmark name/docPr descr/title）的非空值。"""
        values = []
        for node in root.iter():
            local = etree.QName(node).localname
            if local not in _EXTENDED_ATTRIBUTE_SURFACE:
                continue
            for attr_local, attr_ns in _EXTENDED_ATTRIBUTE_SURFACE[local]:
                value = cls._get_extended_attr(node, attr_local, attr_ns)
                if value:
                    values.append(value)
        return values

    def _mask_extended_surface_root(self, root, masker: Masker) -> tuple[int, dict]:
        """对单个 XML root 内的扩展隐私表面执行脱敏。

        - w:instrText、SDT tag、docPr descr/title 直接替换；
        - bookmarkStart/@w:name 替换后校验 XML Name 规范，不符合则 fail closed。
        """
        total = 0
        hits: dict[str, int] = {}

        # 扩展文本节点
        for node in root.iter():
            if etree.QName(node).localname not in _EXTENDED_TEXT_LOCALNAMES:
                continue
            if not node.text:
                continue
            masked, count, chunk_hits = masker.mask_text(node.text)
            if count:
                node.text = masked
                total += count
                for k, v in chunk_hits.items():
                    hits[k] = hits.get(k, 0) + v

        # 扩展属性节点
        for node in root.iter():
            local = etree.QName(node).localname
            if local not in _EXTENDED_ATTRIBUTE_SURFACE:
                continue
            for attr_local, attr_ns in _EXTENDED_ATTRIBUTE_SURFACE[local]:
                value = self._get_extended_attr(node, attr_local, attr_ns)
                if not value:
                    continue
                masked, count, chunk_hits = masker.mask_text(value)
                if not count:
                    continue
                # 书签名称有 XML Name 命名约束，脱敏词不符合则 fail closed
                if local == "bookmarkStart" and not _BOOKMARK_NAME_RE.match(masked):
                    raise RuntimeError(
                        "书签名称包含敏感内容，但脱敏词不符合 XML 名称规范"
                        "（须以字母/下划线开头，仅含字母、数字、下划线、连字符、句点），"
                        "已阻止输出以避免泄漏。请更换脱敏词或从文档中删除该书签。"
                    )
                self._set_extended_attr(node, attr_local, attr_ns, masked)
                total += count
                for k, v in chunk_hits.items():
                    hits[k] = hits.get(k, 0) + v

        return total, hits

    def _restore_extended_surface_root(self, root, restorer: Restorer) -> int:
        """对单个 XML root 内的扩展隐私表面执行恢复（与 _mask_extended_surface_root 对称）。"""
        total = 0
        for node in root.iter():
            if etree.QName(node).localname in _EXTENDED_TEXT_LOCALNAMES and node.text:
                restored, count = restorer.restore_text(node.text)
                if count:
                    node.text = restored
                    total += count
        for node in root.iter():
            local = etree.QName(node).localname
            if local not in _EXTENDED_ATTRIBUTE_SURFACE:
                continue
            for attr_local, attr_ns in _EXTENDED_ATTRIBUTE_SURFACE[local]:
                value = self._get_extended_attr(node, attr_local, attr_ns)
                if not value:
                    continue
                restored, count = restorer.restore_text(value)
                if count:
                    self._set_extended_attr(node, attr_local, attr_ns, restored)
                    total += count
        return total

    def _mask_extended_surface(self, doc: Document, masker: Masker) -> tuple[int, dict]:
        """脱敏所有支持部件内的扩展隐私表面（域代码/SDT 标签/书签名/图片替代文本）。"""
        total = 0
        hits: dict[str, int] = {}

        def _acc(added: int, chunk_hits: dict[str, int]):
            nonlocal total
            total += added
            for k, v in chunk_hits.items():
                hits[k] = hits.get(k, 0) + v

        # document.xml（element 形式，save 时自动序列化）
        _acc(*self._mask_extended_surface_root(doc.element, masker))

        # header/footer（element 形式）
        for root in self._header_footer_roots(doc):
            _acc(*self._mask_extended_surface_root(root, masker))

        # footnotes/endnotes（blob 形式，需写回）
        for _label, note_part in self._note_parts(doc):
            try:
                root = etree.fromstring(note_part.blob)
            except (etree.XMLSyntaxError, ValueError):
                continue
            added, chunk_hits = self._mask_extended_surface_root(root, masker)
            if added:
                note_part._blob = etree.tostring(
                    root, xml_declaration=True, encoding="UTF-8", standalone=True
                )
                _acc(added, chunk_hits)

        # comments（blob 形式，需写回）
        for part in doc.part.package.parts:
            partname = str(part.partname).lstrip("/")
            if not (partname.startswith("word/comments") and partname.endswith(".xml")):
                continue
            try:
                root = etree.fromstring(part.blob)
            except (etree.XMLSyntaxError, ValueError):
                continue
            added, chunk_hits = self._mask_extended_surface_root(root, masker)
            if added:
                self._replace_part_xml(part, root)
                _acc(added, chunk_hits)

        return total, hits

    def _restore_extended_surface(self, doc: Document, restorer: Restorer) -> int:
        """恢复所有支持部件内的扩展隐私表面（与 _mask_extended_surface 对称）。"""
        total = 0

        total += self._restore_extended_surface_root(doc.element, restorer)

        for root in self._header_footer_roots(doc):
            total += self._restore_extended_surface_root(root, restorer)

        for _label, note_part in self._note_parts(doc):
            try:
                root = etree.fromstring(note_part.blob)
            except (etree.XMLSyntaxError, ValueError):
                continue
            added = self._restore_extended_surface_root(root, restorer)
            if added:
                note_part._blob = etree.tostring(
                    root, xml_declaration=True, encoding="UTF-8", standalone=True
                )
                total += added

        for part in doc.part.package.parts:
            partname = str(part.partname).lstrip("/")
            if not (partname.startswith("word/comments") and partname.endswith(".xml")):
                continue
            try:
                root = etree.fromstring(part.blob)
            except (etree.XMLSyntaxError, ValueError):
                continue
            added = self._restore_extended_surface_root(root, restorer)
            if added:
                self._replace_part_xml(part, root)
                total += added

        return total

    @staticmethod
    def _contains_rule_source(text: str, masker: Masker) -> bool:
        if any(key in text for key in masker.codebook.get_sorted_keys()):
            return True
        # A-04: 正则搜索统一走 codebook.safe_search，受编译时 timeout 保护。
        return masker.codebook.safe_search(text)

    @staticmethod
    def _snapshot_all_texts(doc: Document) -> set[str]:
        """脱敏前收集所有 XML 部件的文本内容，用于校验时区分原文和脱敏后文本。

        校验阶段（_assert_no_supported_residuals / _scan_unsupported_parts）
        只对"在快照中存在的文本"检查规则原文残留，从而避免脱敏词自身包含
        规则原文字符（如 {数字占位符1} 中的 '1'）导致的误报。
        """
        texts: set[str] = set()
        for part in doc.part.package.parts:
            partname = str(part.partname).lstrip("/")
            if not partname.endswith(".xml"):
                continue
            try:
                root = etree.fromstring(part.blob)
                for value in root.itertext():
                    if value:
                        texts.add(value)
                # 批注的 author 和 initials 是属性，itertext 不覆盖
                if partname.startswith("word/comments"):
                    for comment in root.findall(".//w:comment", NSMAP):
                        for attribute in (qn("w:author"), qn("w:initials")):
                            value = comment.get(attribute)
                            if value:
                                texts.add(value)
                # 扩展隐私表面的属性值（SDT tag/bookmark name/docPr descr/title）
                # itertext 不覆盖属性，需单独采集
                texts.update(DocxHandler._iter_extended_attr_values(root))
            except (etree.XMLSyntaxError, ValueError):
                continue
        return texts

    def _scan_unsupported_parts(
        self, doc: Document, masker: Masker, original_texts: set[str]
    ) -> list[str]:
        """扫描不支持部件与外部关系，命中时给出不泄露内容的告警。

        与 _assert_no_supported_residuals 同理，只对"在脱敏前快照中存在的文本"
        检查规则原文，避免脱敏词自污染导致的误告警。

        A-03: 除逐节点检查外，增加有界滑动窗口聚合扫描，检测被拆分到相邻
        XML 文本节点中的敏感原文（如 <a:t>SEC</a:t><a:t>RET</a:t>）。
        """
        warnings: list[str] = []
        for part in doc.part.package.parts:
            partname = str(part.partname).lstrip("/")
            if not self._is_supported_part(partname) and partname.endswith(".xml"):
                try:
                    root = etree.fromstring(part.blob)
                    # 按文档顺序收集非空文本值
                    text_values = [v for v in root.itertext() if v]
                    found = False
                    # 1. 逐个文本节点检查（原有逻辑）
                    for value in text_values:
                        if (
                            value in original_texts
                            and self._contains_rule_source(value, masker)
                        ):
                            warnings.append(
                                f"不支持的 OPC 文本部件可能含敏感内容：{partname}"
                            )
                            found = True
                            break
                    # 2. A-03: 有界滑动窗口聚合扫描（检测跨节点拆分）
                    if not found:
                        found = self._scan_cross_node_secrets(
                            text_values, original_texts, masker, partname, warnings
                        )
                except (etree.XMLSyntaxError, ValueError):
                    continue

            for relationship in getattr(part, "rels", {}).values():
                if not relationship.is_external:
                    continue
                target = relationship.target_ref
                if self._contains_rule_source(target, masker):
                    # 安全模式：阻止输出以避免泄漏；兼容模式：仅告警
                    if self.strict_external_targets:
                        raise RuntimeError(
                            f"外部关系目标含敏感内容，已阻止输出以避免泄漏：{partname}"
                        )
                    warnings.append(f"外部关系目标可能含敏感内容：{partname}")
        return list(dict.fromkeys(warnings))

    # A-03: 跨节点拆分敏感原文的滑动窗口扫描上限。
    # 窗口过大会跨不相关结构产生误报；2-4 足以覆盖 OOXML 中常见的 run 拆分。
    _CROSS_NODE_WINDOW_MAX = 4

    def _scan_cross_node_secrets(
        self,
        text_values: list[str],
        original_texts: set[str],
        masker: Masker,
        partname: str,
        warnings: list[str],
    ) -> bool:
        """有界滑动窗口聚合扫描：检测被拆分到相邻文本节点的敏感原文。

        只做告警（warn-only），窗口有界以避免跨不相关结构产生大量误报。
        返回 True 表示已追加告警。
        """
        n = len(text_values)
        for i in range(n):
            for w in range(2, min(self._CROSS_NODE_WINDOW_MAX + 1, n - i + 1)):
                combined = "".join(text_values[i:i + w])
                # A-13: 直接对相邻节点窗口执行规则检测，不再要求 combined
                # 存在于 original_texts（该集合只含逐节点文本，导致 3/4 节点
                # 窗口实际失效）。窗口有界以限制跨不相关结构的误报。
                if self._contains_rule_source(combined, masker):
                    warnings.append(
                        f"不支持的 OPC 文本部件可能含跨节点敏感内容：{partname}"
                    )
                    return True
        return False

    def _assert_no_supported_residuals(
        self, filepath: str, masker: Masker, original_texts: set[str]
    ) -> None:
        """保存后重新读取 ZIP；支持范围内仍有待脱敏内容则拒绝提交。

        A-06: 不再仅检查"在脱敏前快照中存在的文本"，而是检查所有文本值。
        通过排除脱敏词本身来避免误报：如果文本值恰好是某个脱敏词，说明
        该节点已被完全替换，不算残留。部分替换的节点（既不在快照中、
        也不等于任何脱敏词）如果仍包含规则原文，则判定为残留。
        """
        # 收集所有脱敏词，用于排除完全替换的节点
        replacement_words: set[str] = set()
        replacement_words.update(masker.codebook.forward_map.values())
        for _, replacement in masker.codebook.regex_rules:
            if replacement:
                replacement_words.add(replacement)

        with zipfile.ZipFile(filepath) as archive:
            for partname in archive.namelist():
                if not self._is_supported_part(partname) or not partname.endswith(".xml"):
                    continue
                try:
                    root = etree.fromstring(archive.read(partname))
                except etree.XMLSyntaxError as exc:
                    raise RuntimeError(f"保存后 OPC XML 无效：{partname}") from exc

                if partname == "docProps/app.xml" or partname == "docProps/custom.xml":
                    nodes = self._auxiliary_part_nodes(root, partname)
                elif partname == "docProps/core.xml":
                    # 日期、revision 等语义字段不能按普通文本脱敏，否则会破坏类型。
                    nodes = [
                        node for node in root.iter()
                        if etree.QName(node).localname in self.CORE_TEXT_FIELDS
                    ]
                else:
                    # 支持 body 部件：w:t/w:delText + 扩展文本节点（w:instrText）
                    text_localnames = {"t", "delText"} | _EXTENDED_TEXT_LOCALNAMES
                    nodes = [
                        node for node in root.iter()
                        if etree.QName(node).localname in text_localnames
                    ]
                values = [node.text for node in nodes if node.text]
                if partname.startswith("word/comments"):
                    for comment in root.findall(".//w:comment", NSMAP):
                        values.extend(
                            value for value in (
                                comment.get(qn("w:author")),
                                comment.get(qn("w:initials")),
                            ) if value
                        )
                # 扩展隐私表面属性值（SDT tag/bookmark name/docPr descr/title）
                # 必须与写入阶段覆盖完全一致
                if partname not in {"docProps/app.xml", "docProps/custom.xml", "docProps/core.xml"}:
                    values.extend(self._iter_extended_attr_values(root))
                # A-06: 检查所有文本值，排除完全替换的节点（值恰为脱敏词）。
                # 部分替换的节点（值不同于任何脱敏词）仍包含规则原文 -> 残留。
                for value in values:
                    if value in replacement_words:
                        continue
                    if self._contains_rule_source(value, masker):
                        raise RuntimeError(
                            f"保存后残留校验失败，支持部件仍存在待脱敏内容：{partname}"
                        )

    def _assert_no_masked_word_residuals(self, filepath: str, restorer: Restorer) -> None:
        """A-17: 恢复后残留校验，与脱敏阶段对称。

        检查保存后的文件中是否仍有精确规则脱敏词残留。
        正则替换词不可逆，不检查。
        """
        masked_words = set(restorer.codebook.forward_map.values())
        if not masked_words:
            return

        with zipfile.ZipFile(filepath) as archive:
            for partname in archive.namelist():
                if not self._is_supported_part(partname) or not partname.endswith(".xml"):
                    continue
                try:
                    root = etree.fromstring(archive.read(partname))
                except etree.XMLSyntaxError:
                    continue

                if partname == "docProps/app.xml" or partname == "docProps/custom.xml":
                    nodes = self._auxiliary_part_nodes(root, partname)
                elif partname == "docProps/core.xml":
                    nodes = [
                        node for node in root.iter()
                        if etree.QName(node).localname in self.CORE_TEXT_FIELDS
                    ]
                else:
                    text_localnames = {"t", "delText"} | _EXTENDED_TEXT_LOCALNAMES
                    nodes = [
                        node for node in root.iter()
                        if etree.QName(node).localname in text_localnames
                    ]
                values = [node.text for node in nodes if node.text]
                if partname.startswith("word/comments"):
                    for comment in root.findall(".//w:comment", NSMAP):
                        values.extend(
                            v for v in (
                                comment.get(qn("w:author")),
                                comment.get(qn("w:initials")),
                            ) if v
                        )
                if partname not in {"docProps/app.xml", "docProps/custom.xml", "docProps/core.xml"}:
                    values.extend(self._iter_extended_attr_values(root))

                for value in values:
                    if any(word in value for word in masked_words):
                        raise RuntimeError(
                            f"恢复后残留校验失败，支持部件仍存在脱敏词：{partname}"
                        )

    # ======================== 恢复方法 ========================

    def _restore_paragraphs(self, paragraphs, restorer: Restorer) -> int:
        """恢复段落列表"""
        total = 0
        for para in paragraphs:
            total += self._restore_direct_runs(para._p, restorer)
        return total

    def _restore_tables(self, doc: Document, restorer: Restorer) -> int:
        """恢复表格中的文本"""
        total = 0
        def walk(table):
            nonlocal total
            for row in table.rows:
                for cell in row.cells:
                    total += self._restore_paragraphs(cell.paragraphs, restorer)
                    for nested in cell.tables:
                        walk(nested)
                    # 补齐 cell 内的 sdt
                    for sdt_elem in cell._tc.findall(qn("w:sdt")):
                        total += self._restore_sdt_content(sdt_elem, restorer)

        for table in doc.tables:
            walk(table)
        return total

    def _restore_sdt_blocks(self, doc: Document, restorer: Restorer) -> int:
        """恢复 body 中 w:sdt 容器内的段落和表格（与 _mask_sdt_blocks 对应）。"""
        total = 0
        body = doc.element.body
        for sdt_elem in body.findall(qn("w:sdt")):
            total += self._restore_sdt_content(sdt_elem, restorer)
        return total

    def _restore_sdt_content(self, sdt_elem, restorer: Restorer) -> int:
        """递归恢复 w:sdt 内的段落、表格和嵌套 sdt。"""
        total = 0
        content = sdt_elem.find(qn("w:sdtContent"))
        if content is None:
            return 0
        for child in content:
            if child.tag == qn("w:p"):
                total += self._restore_direct_runs(child, restorer)
            elif child.tag == qn("w:tbl"):
                total += self._restore_table_xml(child, restorer)
            elif child.tag == qn("w:sdt"):
                total += self._restore_sdt_content(child, restorer)
        return total

    def _restore_table_xml(self, table_elem, restorer: Restorer) -> int:
        """恢复 XML 表格元素中的文本（递归嵌套表格和 sdt）。"""
        total = 0
        for row in table_elem.findall(qn("w:tr")):
            for cell in row.findall(qn("w:tc")):
                for para_elem in cell.findall(qn("w:p")):
                    total += self._restore_direct_runs(para_elem, restorer)
                for nested_table in cell.findall(qn("w:tbl")):
                    total += self._restore_table_xml(nested_table, restorer)
                for nested_sdt in cell.findall(qn("w:sdt")):
                    total += self._restore_sdt_content(nested_sdt, restorer)
        return total

    def _restore_headers_footers(self, doc: Document, restorer: Restorer) -> int:
        """恢复页眉页脚中的文本"""
        total = 0
        for root in self._header_footer_roots(doc):
            total += self._restore_xml_paragraphs(root, restorer)
        return total

    def _restore_textboxes(self, doc: Document, restorer: Restorer) -> int:
        """恢复文本框中的文本"""
        total = 0
        body = doc.element.body
        txbx_elements = body.findall(".//w:txbxContent", NSMAP)
        for txbx in txbx_elements:
            total += self._restore_xml_paragraphs(txbx, restorer)
        return total

    def _restore_xml_paragraphs(
        self,
        parent_element,
        restorer: Restorer,
        include_hyperlinks: bool = False,
    ) -> int:
        """对 XML 元素中的段落进行恢复"""
        total = 0
        para_elements = parent_element.findall(".//w:p", NSMAP)
        for para_elem in para_elements:
            total += self._restore_direct_runs(para_elem, restorer)
            if include_hyperlinks:
                for hyperlink in para_elem.findall("w:hyperlink", NSMAP):
                    total += self._restore_direct_runs(hyperlink, restorer)
        return total

    def _restore_footnotes(self, doc: Document, restorer: Restorer) -> int:
        """恢复脚注/尾注中的文本"""
        total = 0
        try:
            for _label, note_part in self._note_parts(doc):
                root = etree.fromstring(note_part.blob)
                total += self._restore_xml_paragraphs(
                    root, restorer, include_hyperlinks=True,
                )
                note_part._blob = etree.tostring(
                    root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )
        except Exception as e:
            raise RuntimeError(f"脚注/尾注恢复失败，任务已停止：{e}") from e
        return total

    def _restore_hyperlinks(self, doc: Document, restorer: Restorer) -> int:
        """恢复超链接显示文本"""
        total = 0
        roots = [doc.element.body, *self._header_footer_roots(doc)]
        for root in roots:
            for hyperlink in root.findall(".//w:hyperlink", NSMAP):
                total += self._restore_direct_runs(hyperlink, restorer)
        return total

    def _restore_metadata(self, doc: Document, restorer: Restorer) -> int:
        """恢复文档元数据"""
        total = 0
        props = doc.core_properties
        field_names = (
            "author", "last_modified_by", "title", "subject", "category",
            "comments", "content_status", "identifier", "keywords",
            "language", "version",
        )
        fields = {name: getattr(props, name, None) for name in field_names}
        for field_name, value in fields.items():
            if value:
                restored_text, count = restorer.restore_text(value)
                if count > 0:
                    total += count
                    setattr(props, field_name, restored_text)
        return total

    def _restore_auxiliary_parts(self, doc: Document, restorer: Restorer) -> int:
        """恢复批注、扩展属性和自定义属性 OPC part。"""
        total = 0
        for part in doc.part.package.parts:
            partname = str(part.partname).lstrip("/")
            is_comments = partname.startswith("word/comments") and partname.endswith(".xml")
            if not is_comments and partname not in {"docProps/app.xml", "docProps/custom.xml"}:
                continue
            try:
                root = etree.fromstring(part.blob)
                if is_comments:
                    count = self._restore_xml_paragraphs(
                        root, restorer, include_hyperlinks=True
                    )
                    for comment in root.findall(".//w:comment", NSMAP):
                        for attribute in (qn("w:author"), qn("w:initials")):
                            value = comment.get(attribute)
                            if not value:
                                continue
                            restored, added = restorer.restore_text(value)
                            if added:
                                comment.set(attribute, restored)
                                count += added
                else:
                    count = 0
                    for node in self._auxiliary_part_nodes(root, partname):
                        if not node.text:
                            continue
                        restored, added = restorer.restore_text(node.text)
                        if added:
                            node.text = restored
                            count += added
                if count:
                    self._replace_part_xml(part, root)
                    total += count
            except (etree.XMLSyntaxError, ValueError) as exc:
                raise RuntimeError(f"OPC 部件恢复失败：{partname}") from exc
        return total
